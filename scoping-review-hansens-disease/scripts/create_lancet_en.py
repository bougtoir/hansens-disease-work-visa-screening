#!/usr/bin/env python3
"""Generate Lancet Global Health Comment paper (English) as DOCX with color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

FIGURES_DIR = "/home/ubuntu/scoping_review/figures"
OUTPUT_DIR = "/home/ubuntu/scoping_review/docx"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def create_lancet_comment_en():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    
    # ---- TITLE PAGE ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Comment')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x87, 0x17, 0x17)  # Lancet red
    run.bold = True
    
    doc.add_paragraph()  # blank line
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Hansen's disease remains a barrier to entry: 20 countries still screen work visa applicants for a curable, low-transmission disease")
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()  # blank line
    
    # Authors
    p = doc.add_paragraph()
    run = p.add_run('[Author Name(s)]')
    run.font.size = Pt(12)
    run.italic = True
    
    # Affiliations
    p = doc.add_paragraph()
    run = p.add_run('[Affiliation(s)]')
    run.font.size = Pt(10)
    run.italic = True
    
    # Correspondence
    p = doc.add_paragraph()
    run = p.add_run('Correspondence to: [Corresponding Author Email]')
    run.font.size = Pt(10)
    run.italic = True
    
    doc.add_page_break()
    
    # ---- MAIN TEXT ----
    # Word count note
    p = doc.add_paragraph()
    run = p.add_run('Word count: ~1,500 words | 4 Figures | 1 Table | 15 References')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    # Body text - Paragraph 1
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "Hansen\u2019s disease (leprosy) is among the oldest documented infectious diseases in human history, "
        "yet its place in modern immigration law remains strikingly anachronistic. Despite being curable with "
        "WHO\u2019s multidrug therapy (MDT) within 6\u201312 months, non-highly transmissible (approximately 95% of "
        "the global population has natural immunity to "
    )
    run = p.add_run('Mycobacterium leprae')
    run.italic = True
    p.add_run(
        "), and having a global new case detection rate below 200,000 per year, Hansen\u2019s disease continues "
        "to be listed as a ground for medical inadmissibility in immigration regulations across multiple countries."
    )
    
    # Paragraph 2
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "We conducted a scoping review following PRISMA-ScR guidelines to systematically map which of the "
        "197 UN member states and observer entities (193 member states plus Taiwan, Palestine, Vatican City, "
        "and Kosovo) require disease-specific medical screening for work visa applicants, with particular "
        "attention to whether Hansen\u2019s disease is explicitly named. Our search encompassed official government "
        "legislation, immigration medical examination forms, published regulatory frameworks, the International "
        "Federation of Anti-Leprosy Associations (ILEP) discriminatory laws database, and peer-reviewed academic literature."
    )
    
    # Key Finding paragraph
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run("We identified 20 countries and territories (10\u00b72% of those examined) ")
    run.bold = True
    p.add_run(
        "that explicitly name Hansen\u2019s disease in their work visa medical requirements (Table 1). "
        "The six Gulf Cooperation Council (GCC) states\u2014Saudi Arabia, UAE, Qatar, Kuwait, Oman, and "
        "Bahrain\u2014operate a unified medical screening system through the GCC Approved Medical Centers "
        "Association (GAMCA), which lists leprosy as a condition rendering applicants \u201cmedically unfit.\u201d "
        "Given that the GCC collectively hosts over 35 million migrant workers, predominantly from South "
        "and Southeast Asia, this represents the single largest population affected by Hansen\u2019s disease-based "
        "immigration screening globally."
    )
    
    # Regional patterns
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "Beyond the GCC, five countries in East and Southeast Asia (China, Thailand, Taiwan, Malaysia, and "
        "the Philippines) require certification of non-infection with Hansen\u2019s disease for foreign workers. "
        "South Africa and Namibia retain colonial-era provisions in their immigration medical certificates. "
        "Russia\u2019s Federal Law on the Legal Status of Foreign Citizens includes leprosy among conditions "
        "warranting denial of entry. The United States, despite its 2010 removal of HIV from its list of "
        "inadmissible conditions, continues to classify \u201cinfectious leprosy\u201d as a Class A condition under "
        "INA \u00a7212(a)(1)(A)(i), though in practice the CDC Technical Instructions now allow admission of "
        "individuals who have completed or are undergoing treatment."
    )
    
    # Data accessibility breakdown
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "Our data accessibility varied by language of source. Of 197 countries examined, 110 (55\u00b78%) were "
        "fully characterized through English-language sources alone: 58 with confirmed disease-specific "
        "screening and 52 with confirmed absence of such requirements. For the remaining 87 countries "
        "(44\u00b72%), English sources were insufficient. Supplementary multilingual research in Arabic, "
        "Vietnamese, Sinhala, Indonesian, and other official languages yielded specific screening data for "
        "5 additional countries (2\u00b75%)\u2014Jordan, Lebanon, Vietnam, Sri Lanka, and Indonesia\u2014bringing "
        "confirmed information to 115 countries (58\u00b74%). The remaining 82 countries (41\u00b76%) could not be "
        "fully characterized despite multilingual research; notably, 39 of these have English as an "
        "official language, suggesting non-public dissemination to panel physicians rather than language "
        "barriers as the primary factor. None of the 5 additionally confirmed countries screened for "
        "Hansen\u2019s disease, suggesting that the 20 countries identified may represent a near-complete enumeration."
    )

    # Evidence-policy gap
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "The persistence of Hansen\u2019s disease screening in immigration medical examinations is difficult to "
        "justify on public health grounds. "
    )
    run = p.add_run(
        "Hansen\u2019s disease ranked as the 5th most commonly named disease in work visa screening "
        "(34\u00b75% of 58 countries with disease-specific requirements), despite its low transmissibility, "
        "long incubation period (3\u20135 years), and the availability of free, curative treatment worldwide. "
    )
    run.bold = True
    p.add_run(
        "This positions it alongside diseases with far higher transmission potential and public health "
        "burden. Tuberculosis (94\u00b78%), HIV/AIDS (82\u00b78%), syphilis (65\u00b75%), and hepatitis B (51\u00b77%) "
        "are all more frequently screened, but each carries demonstrably greater epidemiological justification "
        "for pre-entry detection."
    )
    
    # Japan comparison
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "Japan\u2019s experience offers an instructive counter-example. Having abolished its Leprosy Prevention "
        "Law (\u3089\u3044\u4e88\u9632\u6cd5) in 1996 and acknowledged through landmark litigation in 2001 that forced "
        "isolation constituted a violation of constitutional rights, Japan does not include Hansen\u2019s disease "
        "in its immigration medical screening. Japan\u2019s pre-entry TB screening programme (JPETS) demonstrates "
        "that evidence-based, disease-specific screening is feasible without perpetuating stigma against "
        "conditions that do not meet the threshold for border health intervention."
    )
    
    # UN Resolution
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "The UN General Assembly Resolution A/RES/65/215 (2010) calls upon member states to eliminate "
        "discrimination against persons affected by Hansen\u2019s disease, with Principle 7 explicitly affirming "
        "the right of affected persons to freedom of movement, including entry into and exit from countries. "
        "The UN Special Rapporteur on leprosy, appointed in 2017, has documented ongoing discriminatory laws "
        "in immigration contexts. Yet our findings indicate that at least 20 countries maintain provisions "
        "that directly contravene these principles."
    )
    
    # Call to action
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "We call on the following actions: (1) Countries currently screening for Hansen\u2019s disease in "
        "immigration medical examinations should review these requirements against current epidemiological "
        "evidence and WHO guidelines; (2) the GCC states should consider removing leprosy from the "
        "GAMCA/WAFID medical unfitness criteria, given the disproportionate impact on millions of migrant "
        "workers; (3) WHO should issue explicit guidance that Hansen\u2019s disease does not meet the threshold "
        "for pre-entry screening at international borders; and (4) the UN Human Rights Council should "
        "monitor progress in eliminating Hansen\u2019s disease-based immigration restrictions as part of the "
        "Special Rapporteur\u2019s mandate."
    )
    
    # Closing
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "The administrative requirement to certify absence of Hansen\u2019s disease for work visa purposes "
        "represents a clear divergence from standard medical practice (\u6a19\u6e96\u533b\u7642\u3068\u306e\u4e56\u96e2). "
        "When a disease is curable, non-highly transmissible, and undetectable through routine examination "
        "due to its long incubation period, border screening serves neither public health nor individual "
        "welfare\u2014it serves only stigma. The time to align immigration medical requirements with modern "
        "evidence is long overdue."
    )
    
    doc.add_page_break()
    
    # ---- TABLE 1 ----
    p = doc.add_paragraph()
    run = p.add_run("Table 1: Countries and territories explicitly naming Hansen\u2019s disease in work visa medical requirements")
    run.bold = True
    run.font.size = Pt(10)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header
    headers = ['Region', 'Country/Territory', 'Legal Instrument', 'Provision Type']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'D32F2F')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Data
    data = [
        ['GCC / Middle East', 'Saudi Arabia', 'GAMCA/WAFID Regulations (6th Version, 2021)', 'Automatic exclusion'],
        ['', 'UAE', 'GAMCA/WAFID Regulations', 'Automatic exclusion'],
        ['', 'Qatar', 'GAMCA/WAFID Regulations', 'Automatic exclusion'],
        ['', 'Kuwait', 'GAMCA/WAFID Regulations', 'Automatic exclusion'],
        ['', 'Oman', 'GAMCA/WAFID Regulations', 'Automatic exclusion'],
        ['', 'Bahrain', 'GAMCA/WAFID Regulations', 'Automatic exclusion'],
        ['East / SE Asia', 'China', 'Foreigner Physical Examination Form', 'Certification of absence'],
        ['', 'Thailand', 'Emergency Decree on Managing Foreign Workers B.E. 2560', 'Prohibited disease'],
        ['', 'Taiwan', 'Employment Gold Card / Work Permit Regulations', 'Certification of absence'],
        ['', 'Malaysia', 'FOMEMA Medical Examination Requirements', 'Medical unfitness'],
        ['', 'Philippines', 'Overseas Workers Welfare Admin. Requirements', 'Certification of absence'],
        ['Africa', 'South Africa', 'Immigration Act 2002; Form BI-811', 'Certification of absence'],
        ['', 'Namibia', 'Immigration Control Act 1993', 'Medical prohibition'],
        ['Europe', 'Russia', 'Federal Law No. 115-FZ (2002)', 'Deportation ground'],
        ['', 'Malta', 'Immigration Act (Cap. 217)', 'Medical inadmissibility'],
        ['Americas', 'United States', 'INA \u00a7212(a)(1)(A)(i); 42 CFR \u00a734.2', 'Class A condition*'],
        ['', 'Barbados', 'Immigration Act (Cap. 190)', 'Medical prohibition'],
        ['', 'US Virgin Islands', 'Follows US federal immigration law', 'Class A condition*'],
        ['South Asia', 'India', 'Various state employment laws; 108 discriminatory laws (ILEP)', 'Employment restriction'],
        ['', 'Hong Kong SAR', 'Immigration Ordinance (Cap. 115)', 'Medical inadmissibility'],
    ]
    
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    # Table note
    p = doc.add_paragraph()
    run = p.add_run(
        '*Note: USA classifies "infectious leprosy" as Class A but CDC Technical Instructions (2018) '
        'allow waiver for individuals completing or having completed MDT treatment.'
    )
    run.font.size = Pt(8)
    run.italic = True
    
    doc.add_page_break()
    
    # ---- FIGURE ----
    p = doc.add_paragraph()
    run = p.add_run("Figure 1")
    run.bold = True
    run.font.size = Pt(10)
    
    fig_path = os.path.join(FIGURES_DIR, "fig5_transmissibility.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        "Figure 1: Relationship between disease transmissibility and frequency of screening "
        "in work visa medical examinations across 58 countries. Hansen\u2019s disease (red) occupies "
        "a unique position: low transmissibility but high screening frequency, indicating an "
        "evidence-policy gap. Bubble size proportional to number of countries screening for each disease."
    )
    run.font.size = Pt(9)
    run.italic = True
    
    doc.add_page_break()
    
    # ---- FIGURE 2: World Map ----
    p = doc.add_paragraph()
    run = p.add_run("Figure 2")
    run.bold = True
    run.font.size = Pt(10)
    
    fig_path = os.path.join(FIGURES_DIR, "fig1_world_map.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        "Figure 2: Global distribution of Hansen\u2019s disease screening in work visa medical "
        "examinations. Red: Hansen\u2019s disease explicitly named (20 countries); Blue: disease-specific "
        "screening without Hansen\u2019s disease (38 countries); Green: no disease-specific medical exam "
        "required (52+ countries); Gray: information not publicly available."
    )
    run.font.size = Pt(9)
    run.italic = True
    
    doc.add_page_break()

    # ---- FIGURE 3: Sankey Diagram ----
    p = doc.add_paragraph()
    run = p.add_run("Figure 3")
    run.bold = True
    run.font.size = Pt(10)

    fig_path = os.path.join(FIGURES_DIR, "fig6_sankey_en.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "Figure 3: Data accessibility flow diagram showing the two-stage search strategy. "
        "Of 197 countries examined, 110 were characterized via English sources (58 with disease-specific "
        "screening confirmed, 52 with confirmed absence), 5 via multilingual research, and 82 remained "
        "unreachable. Among 63 countries with confirmed screening, 20 explicitly named Hansen\u2019s disease."
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    # ---- FIGURE 4: Accessibility Map ----
    p = doc.add_paragraph()
    run = p.add_run("Figure 4")
    run.bold = True
    run.font.size = Pt(10)

    fig_path = os.path.join(FIGURES_DIR, "fig7_accessibility_en.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "Figure 4: Geographic distribution of data accessibility for work visa medical requirements. "
        "Green: reached via English-language sources (110 countries, 55.8%); Orange: reached via "
        "multilingual research (5 countries, 2.5%); Red: unreachable (82 countries, 41.6%)."
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()
    
    # ---- REFERENCES ----
    p = doc.add_paragraph()
    run = p.add_run('References')
    run.bold = True
    run.font.size = Pt(12)
    
    refs = [
        "1. WHO. Weekly epidemiological record: Global leprosy (Hansen disease) update, 2024. Wkly Epidemiol Rec 2025; 100: 377\u201396.",
        "2. Tricco AC, Lillie E, Zarin W, et al. PRISMA Extension for Scoping Reviews (PRISMA-ScR): Checklist and explanation. Ann Intern Med 2018; 169: 467\u201373.",
        "3. GCC Approved Medical Centers Association. WAFID Medical Examination Regulations, 6th Version. Riyadh: GAMCA Secretariat, 2021.",
        "4. Republic of South Africa. Immigration Act No. 13 of 2002; Form BI-811 (Medical Certificate). Department of Home Affairs.",
        "5. United States Congress. Immigration and Nationality Act \u00a7212(a)(1)(A)(i); 42 CFR \u00a734.2 (Medical examination of aliens).",
        "6. CDC. Technical Instructions for Hansen\u2019s Disease for Panel Physicians. Atlanta: Centers for Disease Control and Prevention, 2018.",
        "7. Kingdom of Thailand. Emergency Decree on Managing the Work of Aliens B.E. 2560 (2017). Royal Thai Government Gazette.",
        "8. People\u2019s Republic of China. Foreigner Physical Examination Form. Ministry of Health.",
        "9. Russian Federation. Federal Law No. 115-FZ on the Legal Status of Foreign Citizens (2002), as amended.",
        "10. ILEP. Discriminatory Laws Database. International Federation of Anti-Leprosy Associations, 2024. https://ilepfederation.org/discriminatory-laws/",
        "11. UN General Assembly. Resolution A/RES/65/215: Elimination of discrimination against persons affected by leprosy and their family members. New York: United Nations, 2010.",
        "12. UN Human Rights Council. Principles and guidelines for the elimination of discrimination against persons affected by leprosy and their family members. A/HRC/RES/15/10, 2010.",
        "13. Nanri T. Initiatives to address leprosy as a human rights issue through the mandate of UN Special Rapporteur. PLoS Negl Trop Dis 2022; 16: e0010201.",
        "14. Abubakar I, Aldridge RW, Devakumar D, et al. The UCL\u2013Lancet Commission on Migration and Health. Lancet 2018; 392: 2606\u201354.",
        "15. Atkins H. Healthy Enough to Enter? Exploring the nexus of the body and the border through South African visa medical requirements. Link\u00f6ping University, 2020.",
    ]
    
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
    
    # ---- DECLARATION ----
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run('Declaration of interests')
    run.bold = True
    p = doc.add_paragraph()
    p.add_run('We declare no competing interests.')
    
    p = doc.add_paragraph()
    run = p.add_run('Acknowledgements')
    run.bold = True
    p = doc.add_paragraph()
    p.add_run('[To be added]')
    
    # Save
    output_path = os.path.join(OUTPUT_DIR, "Lancet_Global_Health_Comment_EN.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == '__main__':
    create_lancet_comment_en()
