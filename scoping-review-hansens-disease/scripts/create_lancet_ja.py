#!/usr/bin/env python3
"""Generate Lancet Global Health Comment paper (Japanese) as DOCX with color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FIGURES_DIR = "/home/ubuntu/scoping_review/figures"
OUTPUT_DIR = "/home/ubuntu/scoping_review/docx"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def create_lancet_comment_ja():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    
    # ---- TITLE PAGE ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Comment（コメント）')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x87, 0x17, 0x17)
    run.bold = True
    
    doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("ハンセン病は依然として入国の障壁である：治癒可能かつ低感染性の疾病に対し、20カ国が労働ビザ申請者のスクリーニングを継続")
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph()
    
    # English subtitle
    p = doc.add_paragraph()
    run = p.add_run("Hansen\u2019s disease remains a barrier to entry: 20 countries still screen work visa applicants for a curable, low-transmission disease")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('[著者名]')
    run.font.size = Pt(12)
    run.italic = True
    
    p = doc.add_paragraph()
    run = p.add_run('[所属機関]')
    run.font.size = Pt(10)
    run.italic = True
    
    p = doc.add_paragraph()
    run = p.add_run('連絡先: [責任著者メールアドレス]')
    run.font.size = Pt(10)
    run.italic = True
    
    doc.add_page_break()
    
    # ---- MAIN TEXT ----
    p = doc.add_paragraph()
    run = p.add_run('語数: 約1,500語（英語換算） | 図4点 | 表1点 | 参考文献15件')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    run = p.add_run('※本文書はLancet Global Health Comment投稿用原稿の日本語版です。投稿時には英語版を使用してください。')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    
    doc.add_paragraph()
    
    # Body - Paragraph 1
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "ハンセン病（Hansen\u2019s disease, leprosy）は人類の歴史上最も古くから記録されている感染症の一つであるが、"
        "現代の移民法における位置づけは驚くほど時代錯誤的である。WHOの多剤併用療法（MDT）により6〜12ヶ月で"
        "完治すること、高度に感染性ではないこと（世界人口の約95%が"
    )
    run = p.add_run('Mycobacterium leprae')
    run.italic = True
    p.add_run(
        "に対する自然免疫を有する）、そして世界の年間新規発見症例数が20万人未満であるにもかかわらず、"
        "ハンセン病は複数の国において移民法上の医療上の入国不許可事由として依然掲載されている。"
    )
    
    # Paragraph 2
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "我々はPRISMA-ScRガイドラインに従いスコーピングレビューを実施し、197の国連加盟国・オブザーバー国"
        "（加盟国193カ国＋台湾、パレスチナ、バチカン市国、コソボ）のうち、労働ビザ申請者に対して疾病特異的な"
        "医療検査を要求している国を体系的にマッピングし、特にハンセン病が明示的に記載されているかに注目した。"
        "調査対象は、各国の公式政府法令、移民医療検査書式、公表された規制枠組み、国際ハンセン病団体連合（ILEP）"
        "の差別的法律データベース、および査読付き学術文献を包含した。"
    )
    
    # Key finding
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run("我々は20カ国・地域（調査対象の10.2%）がハンセン病を労働ビザ医療要件に明示的に記載していることを特定した")
    run.bold = True
    p.add_run(
        "（表1）。湾岸協力会議（GCC）6カ国—サウジアラビア、UAE、カタール、クウェート、オマーン、バーレーン—"
        "はGCC認定医療センター協会（GAMCA）を通じた統一医療検査制度を運用しており、ハンセン病を「医療上不適格」"
        "とする条件として記載している。GCC諸国は南・東南アジア出身者を中心に合計3,500万人以上の外国人労働者を"
        "受け入れており、これはハンセン病に基づく移民医療検査として世界最大の影響を持つ制度である。"
    )
    
    # Regional patterns
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "GCC以外では、東・東南アジア5カ国・地域（中国、タイ、台湾、マレーシア、フィリピン）が外国人労働者に"
        "対してハンセン病の不罹患証明を要求している。南アフリカとナミビアは植民地時代に起源を持つ規定を移民"
        "医療証明書に残存させている。ロシア連邦の「外国人の法的地位に関する連邦法」はハンセン病を入国拒否事由"
        "に含めている。米国は2010年にHIVを入国不許可条件リストから削除したにもかかわらず、移民国籍法（INA）"
        "§212(a)(1)(A)(i)に基づき「伝染性ハンセン病」をClass A条件として分類し続けている。ただし実務上は、"
        "CDCの技術指示書（2018年）により治療中または治療完了者の入国は認められている。"
    )
    
    # Data accessibility breakdown
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "データへの到達は情報源の言語によって異なった。調査した197カ国のうち、110カ国（55.8%）は英語情報源のみで"
        "完全に特定された：疾病特異的スクリーニング確認58カ国、スクリーニング不存在確認52カ国。残りの87カ国"
        "（44.2%）は英語情報源では不十分であった。アラビア語、ベトナム語、シンハラ語、インドネシア語等の"
        "公用語による補足的多言語調査により、5カ国（2.5%）—ヨルダン、レバノン、ベトナム、スリランカ、"
        "インドネシア—の具体的なスクリーニングデータが得られ、確認済み情報は115カ国（58.4%）となった。"
        "残りの82カ国（41.6%）は多言語調査にもかかわらず完全に特定できなかった。これら82カ国のうち39カ国は"
        "英語を公用語としており、言語障壁よりも指定医への非公開の要件伝達が主要因であることを示唆している。"
        "追加確認された5カ国のいずれもハンセン病をスクリーニングしておらず、特定された20カ国がほぼ完全な"
        "列挙である可能性が示唆された。"
    )

    # Evidence-policy gap
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "移民医療検査におけるハンセン病スクリーニングの継続を公衆衛生上の根拠から正当化することは困難である。"
    )
    run = p.add_run(
        "ハンセン病は労働ビザスクリーニングにおいて5番目に多く指定される疾病（疾病特異的要件を持つ58カ国中34.5%）"
        "であったが、その低い感染性、長い潜伏期間（3〜5年）、そして世界中で無償の根治療法が利用可能であることを"
        "考慮すると、これは標準医療との明白な乖離である。"
    )
    run.bold = True
    p.add_run(
        "結核（94.8%）、HIV/AIDS（82.8%）、梅毒（65.5%）、B型肝炎（51.7%）はいずれもより高頻度でスクリーニング"
        "されているが、それぞれ入国前検出に対するより大きな疫学的正当性を有している。"
    )
    
    # Japan comparison
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "日本の経験は示唆的な反例を提供する。1996年の「らい予防法」廃止と2001年の画期的な熊本地裁判決を経て"
        "強制隔離が憲法上の権利侵害であることを認めた日本は、移民医療検査にハンセン病を含めていない。"
        "日本の入国前結核スクリーニングプログラム（JPETS）は、スティグマを永続させることなく、エビデンスに基づく"
        "疾病特異的スクリーニングが実現可能であることを示している。"
    )
    
    # UN Resolution
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "国連総会決議A/RES/65/215（2010年）はハンセン病罹患者に対する差別の撤廃を加盟国に求めており、原則7は"
        "罹患者の移動の自由の権利を明示的に確認している。2017年に任命された国連ハンセン病特別報告者は、"
        "移民文脈における差別的法律を文書化してきた。しかし我々の知見は、少なくとも20カ国がこれらの原則に"
        "直接矛盾する規定を維持していることを示している。"
    )
    
    # Call to action
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "我々は以下の行動を求める：（1）現在移民医療検査でハンセン病をスクリーニングしている国は、これらの要件を"
        "現在の疫学的エビデンスおよびWHOガイドラインに照らして見直すべきである；（2）GCC諸国は、数百万人の"
        "外国人労働者に対する不均衡な影響を考慮し、GAMCA/WAFIDの医療不適格基準からハンセン病を除外することを"
        "検討すべきである；（3）WHOはハンセン病が国際的な国境での入国前スクリーニングの閾値を満たさないことに"
        "ついて明確なガイダンスを発出すべきである；（4）国連人権理事会は、特別報告者の権限の一部として、"
        "ハンセン病に基づく移民制限の撤廃における進捗を監視すべきである。"
    )
    
    # Closing
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        "労働ビザ取得のためにハンセン病の不罹患を証明するという行政要求は、標準医療との明確な乖離を示している。"
        "治癒可能であり、高度に感染性ではなく、長い潜伏期間のために定期的な検査で検出不可能な疾病に対して、"
        "国境でのスクリーニングは公衆衛生上の利益にも個人の福祉にも寄与しない—それはスティグマにのみ寄与する。"
        "移民医療要件を現代のエビデンスに整合させる時期は、とうの昔に過ぎている。"
    )
    
    doc.add_page_break()
    
    # ---- TABLE 1 ----
    p = doc.add_paragraph()
    run = p.add_run("表1：労働ビザ医療要件にハンセン病を明示的に記載している国・地域")
    run.bold = True
    run.font.size = Pt(10)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ['地域', '国・地域', '法的根拠', '規定の性質']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'D32F2F')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    data = [
        ['GCC／中東', 'サウジアラビア', 'GAMCA/WAFID規則（第6版、2021年）', '自動排除'],
        ['', 'UAE', 'GAMCA/WAFID規則', '自動排除'],
        ['', 'カタール', 'GAMCA/WAFID規則', '自動排除'],
        ['', 'クウェート', 'GAMCA/WAFID規則', '自動排除'],
        ['', 'オマーン', 'GAMCA/WAFID規則', '自動排除'],
        ['', 'バーレーン', 'GAMCA/WAFID規則', '自動排除'],
        ['東・東南アジア', '中国', '外国人体格検査表', '不罹患証明'],
        ['', 'タイ', '外国人労働管理緊急勅令 B.E.2560', '禁止疾病'],
        ['', '台湾', '就業ゴールドカード／労働許可規則', '不罹患証明'],
        ['', 'マレーシア', 'FOMEMA医療検査要件', '医療上不適格'],
        ['', 'フィリピン', '海外労働者福祉庁（OWWA）要件', '不罹患証明'],
        ['アフリカ', '南アフリカ', '移民法（2002年）；BI-811書式', '不罹患証明'],
        ['', 'ナミビア', '移民管理法（1993年）', '医療上の禁止'],
        ['ヨーロッパ', 'ロシア', '連邦法No.115-FZ（2002年）', '国外退去事由'],
        ['', 'マルタ', '移民法（Cap.217）', '医療上の入国不許可'],
        ['アメリカ大陸', '米国', 'INA §212(a)(1)(A)(i); 42 CFR §34.2', 'Class A条件*'],
        ['', 'バルバドス', '移民法（Cap.190）', '医療上の禁止'],
        ['', '米領ヴァージン諸島', '米国連邦移民法に準拠', 'Class A条件*'],
        ['南アジア', 'インド', '各州雇用法；108の差別的法律（ILEP）', '雇用制限'],
        ['', '香港特別行政区', '入国管理条例（Cap.115）', '医療上の入国不許可'],
    ]
    
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    p = doc.add_paragraph()
    run = p.add_run(
        '*注：米国は「伝染性ハンセン病」をClass Aに分類しているが、CDC技術指示書（2018年）により'
        'MDT治療を完了した、または治療中の者については免除が認められる。'
    )
    run.font.size = Pt(8)
    run.italic = True
    
    doc.add_page_break()
    
    # ---- FIGURE ----
    p = doc.add_paragraph()
    run = p.add_run("図1")
    run.bold = True
    run.font.size = Pt(10)
    
    fig_path = os.path.join(FIGURES_DIR, "fig5_transmissibility.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        "図1：58カ国における労働ビザ医療検査での疾病の感染性とスクリーニング頻度の関係。"
        "ハンセン病（赤）は低い感染性にもかかわらず高いスクリーニング頻度を示し、エビデンスと政策の"
        "乖離を示唆する独特の位置を占める。バブルサイズは各疾病をスクリーニングしている国の数に比例。"
    )
    run.font.size = Pt(9)
    run.italic = True
    
    doc.add_page_break()
    
    # ---- 図2: 世界地図 ----
    p = doc.add_paragraph()
    run = p.add_run("図2")
    run.bold = True
    run.font.size = Pt(10)
    
    fig_path = os.path.join(FIGURES_DIR, "fig1_world_map.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(
        "図2：労働ビザ医療検査におけるハンセン病スクリーニングの世界分布。"
        "赤：ハンセン病を明示的に記載（20カ国）、青：ハンセン病を含まない疾病特異的スクリーニング（38カ国）、"
        "緑：疾病特異的医療検査を要求しない（52カ国以上）、灰色：公開情報なし。"
    )
    run.font.size = Pt(9)
    run.italic = True
    
    doc.add_page_break()

    # ---- 図3: Sankey Diagram ----
    p = doc.add_paragraph()
    run = p.add_run("図3")
    run.bold = True
    run.font.size = Pt(10)

    fig_path = os.path.join(FIGURES_DIR, "fig6_sankey_ja.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "図3：データ到達フロー図（2段階検索戦略）。調査した197カ国のうち、110カ国は英語情報源で"
        "特定（疾病特異的スクリーニング確認58カ国、不存在確認52カ国）、5カ国は多言語調査で特定、"
        "82カ国は到達不可。スクリーニング確認63カ国のうち20カ国がハンセン病を明示的に記載。"
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    # ---- 図4: Accessibility Map ----
    p = doc.add_paragraph()
    run = p.add_run("図4")
    run.bold = True
    run.font.size = Pt(10)

    fig_path = os.path.join(FIGURES_DIR, "fig7_accessibility_ja.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        "図4：労働ビザ医療要件に関するデータ到達可能性の地理的分布。"
        "緑：英語情報源で到達（110カ国、55.8%）、橙：多言語調査で到達（5カ国、2.5%）、"
        "赤：到達不可（82カ国、41.6%）。"
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()
    
    # ---- REFERENCES ----
    p = doc.add_paragraph()
    run = p.add_run('参考文献')
    run.bold = True
    run.font.size = Pt(12)
    
    refs = [
        "1. WHO. Weekly epidemiological record: Global leprosy (Hansen disease) update, 2024. Wkly Epidemiol Rec 2025; 100: 377\u201396.",
        "2. Tricco AC, Lillie E, Zarin W, et al. PRISMA Extension for Scoping Reviews (PRISMA-ScR): Checklist and explanation. Ann Intern Med 2018; 169: 467\u201373.",
        "3. GCC Approved Medical Centers Association. WAFID Medical Examination Regulations, 6th Version. Riyadh: GAMCA Secretariat, 2021.",
        "4. Republic of South Africa. Immigration Act No. 13 of 2002; Form BI-811 (Medical Certificate). Department of Home Affairs.",
        "5. United States Congress. Immigration and Nationality Act \u00a7212(a)(1)(A)(i); 42 CFR \u00a734.2 (Medical examination of aliens).",
        "6. CDC. Technical Instructions for Hansen\u2019s Disease for Panel Physicians. Atlanta: Centers for Disease Control and Prevention, 2018.",
        "7. Kingdom of Thailand. Emergency Decree on Managing the Work of Aliens B.E. 2560 (2017). Royal Thai Government Gazette.",
        "8. People\u2019s Republic of China. Foreigner Physical Examination Form. Ministry of Health.",
        "9. Russian Federation. Federal Law No. 115-FZ on the Legal Status of Foreign Citizens (2002), as amended.",
        "10. ILEP. Discriminatory Laws Database. International Federation of Anti-Leprosy Associations, 2024. https://ilepfederation.org/discriminatory-laws/",
        "11. UN General Assembly. Resolution A/RES/65/215: Elimination of discrimination against persons affected by leprosy and their family members. New York: United Nations, 2010.",
        "12. UN Human Rights Council. Principles and guidelines for the elimination of discrimination against persons affected by leprosy and their family members. A/HRC/RES/15/10, 2010.",
        "13. Nanri T. Initiatives to address leprosy as a human rights issue through the mandate of UN Special Rapporteur. PLoS Negl Trop Dis 2022; 16: e0010201.",
        "14. Abubakar I, Aldridge RW, Devakumar D, et al. The UCL\u2013Lancet Commission on Migration and Health. Lancet 2018; 392: 2606\u201354.",
        "15. Atkins H. Healthy Enough to Enter? Exploring the nexus of the body and the border through South African visa medical requirements. Link\u00f6ping University, 2020.",
    ]
    
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
    
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run('利益相反の開示')
    run.bold = True
    p = doc.add_paragraph()
    p.add_run('利益相反はない。')
    
    p = doc.add_paragraph()
    run = p.add_run('謝辞')
    run.bold = True
    p = doc.add_paragraph()
    p.add_run('[追記予定]')
    
    output_path = os.path.join(OUTPUT_DIR, "Lancet_Global_Health_Comment_JA.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == '__main__':
    create_lancet_comment_ja()
