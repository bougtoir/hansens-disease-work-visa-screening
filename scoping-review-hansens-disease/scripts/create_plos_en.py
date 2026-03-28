#!/usr/bin/env python3
"""Generate PLoS NTDs Full Paper (English) as DOCX with color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FIGURES_DIR = "/home/ubuntu/scoping_review/figures"
OUTPUT_DIR = "/home/ubuntu/scoping_review/docx"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)  # PLoS dark blue
    return h


def add_body(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_table_data(table, data, font_size=8):
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def create_plos_full_en():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    # ==== TITLE PAGE ====
    p = doc.add_paragraph()
    run = p.add_run('RESEARCH ARTICLE')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "Hansen\u2019s Disease Screening in Work Visa Medical Examinations: "
        "A Global Scoping Review of Administrative Requirements That Diverge from Standard Medical Practice"
    )
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('[Author Name(s)]')
    run.font.size = Pt(12)
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('[Affiliation(s)]')
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('* Corresponding author: [email]')
    run.font.size = Pt(10)
    run.italic = True

    doc.add_page_break()

    # ==== ABSTRACT ====
    add_heading_styled(doc, 'Abstract', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Background')
    run.bold = True
    p.add_run(
        '\nHansen\u2019s disease (leprosy) is a curable, low-transmissibility infection caused by '
        'Mycobacterium leprae. Despite advances in treatment and declining global incidence, '
        'some countries continue to require screening for Hansen\u2019s disease as part of work visa '
        'medical examinations. The extent and nature of these requirements have not been systematically mapped.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Objectives')
    run.bold = True
    p.add_run(
        '\nTo systematically identify which countries require disease-specific medical screening '
        'for work visa applicants, with particular attention to Hansen\u2019s disease; to characterize '
        'the nature of these requirements; and to evaluate their alignment with current epidemiological '
        'evidence and international human rights standards.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Methods')
    run.bold = True
    p.add_run(
        '\nWe conducted a scoping review following PRISMA-ScR guidelines across 197 countries and '
        'territories (193 UN member states plus 4 observer entities). Sources included official '
        'government legislation, immigration medical examination forms, regulatory frameworks, the '
        'ILEP discriminatory laws database, IOM reports, and peer-reviewed literature. Data were '
        'extracted on disease-specific screening requirements, legal instruments, and the nature of provisions.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Results')
    run.bold = True
    p.add_run(
        '\nOf 197 countries/territories examined, 58 (29.4%) have confirmed disease-specific medical '
        'screening requirements for work visas. Twenty countries/territories (10.2%) explicitly name '
        'Hansen\u2019s disease. The six GCC states operate a unified GAMCA/WAFID system listing leprosy '
        'as \u201cmedically unfit,\u201d affecting over 35 million migrant workers. Hansen\u2019s disease was the '
        '5th most commonly screened condition (34.5%), after tuberculosis (94.8%), HIV/AIDS (82.8%), '
        'syphilis (65.5%), and hepatitis B (51.7%). Twenty-six EU/Schengen countries and 26 additional '
        'countries require no disease-specific screening. Hansen\u2019s disease screening shows no '
        'correlation with national leprosy burden.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Conclusions')
    run.bold = True
    p.add_run(
        '\nHansen\u2019s disease remains embedded in immigration medical requirements in 20 countries '
        'despite lacking epidemiological justification. These requirements represent administrative '
        'demands that diverge from standard medical practice and contravene UN Resolution A/RES/65/215. '
        'Urgent review and reform of these provisions is warranted.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.bold = True
    p.add_run(
        'Hansen\u2019s disease; leprosy; work visa; immigration medical examination; scoping review; '
        'PRISMA-ScR; medical inadmissibility; discrimination; neglected tropical diseases'
    )

    doc.add_page_break()

    # ==== AUTHOR SUMMARY ====
    add_heading_styled(doc, 'Author Summary', level=1)
    p = doc.add_paragraph()
    p.add_run(
        'Hansen\u2019s disease (leprosy) is one of the oldest known infectious diseases, but it is now '
        'fully curable with WHO\u2019s multidrug therapy and is not highly contagious\u201495% of people have '
        'natural immunity. Despite this, we found that 20 countries still require work visa applicants '
        'to be screened for this disease. The six Gulf Cooperation Council states (Saudi Arabia, UAE, '
        'Qatar, Kuwait, Oman, and Bahrain) collectively host over 35 million foreign workers and list '
        'leprosy as a condition making applicants \u201cmedically unfit.\u201d This is particularly striking '
        'because these requirements persist even though the UN General Assembly passed a resolution in '
        '2010 calling for elimination of discrimination against persons affected by leprosy, including '
        'in immigration contexts. Our study is the first to systematically map these requirements across '
        'all 197 UN-recognized countries and territories. We found that Hansen\u2019s disease is the 5th '
        'most commonly screened condition in work visa medical exams, despite having far lower '
        'transmissibility than the top four (tuberculosis, HIV, syphilis, and hepatitis B). Japan\u2019s '
        'experience\u2014having abolished its Leprosy Prevention Law in 1996 without including leprosy '
        'in immigration screening\u2014demonstrates that evidence-based approaches are feasible. We call '
        'for countries to align their immigration medical requirements with current scientific evidence '
        'and international human rights standards.'
    )

    doc.add_page_break()

    # ==== INTRODUCTION ====
    add_heading_styled(doc, '1. Introduction', level=1)

    add_body(doc,
        'Hansen\u2019s disease (leprosy), caused by Mycobacterium leprae, is among the oldest documented '
        'infectious diseases in human history [1]. Despite dramatic advances in treatment\u2014notably the '
        'introduction of WHO\u2019s multidrug therapy (MDT) in 1981, which renders patients non-infectious '
        'within days and achieves complete cure within 6\u201312 months [2]\u2014the disease continues to carry '
        'profound social stigma in many parts of the world [3].'
    )

    add_body(doc,
        'The global epidemiological picture has improved substantially. The WHO reported approximately '
        '182,000 new cases detected globally in 2024, with new case detection rates declining steadily '
        'over the past two decades [4]. Approximately 95% of the global population possesses natural '
        'immunity to M. leprae, and transmission requires prolonged close contact with untreated '
        'multibacillary patients [5]. The incubation period ranges from 3 to 5 years (and up to 20 '
        'years in some cases), making point-of-entry detection through medical examination largely '
        'ineffective [6].'
    )

    add_body(doc,
        'Despite these characteristics, Hansen\u2019s disease has historically been included in immigration '
        'medical inadmissibility lists in various countries. The persistence of such provisions raises '
        'important questions about the alignment between administrative requirements for entry and '
        'current medical knowledge\u2014what we term the divergence between administrative demands and '
        'standard medical practice (\u6a19\u6e96\u533b\u7642\u3068\u4e56\u96e2\u3057\u305f\u884c\u653f\u8981\u6c42).'
    )

    add_body(doc,
        'The international community has explicitly addressed leprosy-related discrimination. UN General '
        'Assembly Resolution A/RES/65/215 (2010) calls upon states to eliminate discrimination against '
        'persons affected by Hansen\u2019s disease, with Principle 7 affirming the right to freedom of '
        'movement including entry into countries [7]. A UN Special Rapporteur on the elimination of '
        'discrimination against persons affected by leprosy was appointed in 2017 [8]. The International '
        'Federation of Anti-Leprosy Associations (ILEP) maintains a database documenting 139 '
        'discriminatory laws across 24 countries [9].'
    )

    add_body(doc,
        'However, no comprehensive, systematic mapping of Hansen\u2019s disease screening requirements in '
        'work visa medical examinations has been conducted globally. Existing literature addresses '
        'individual countries or regions but lacks a unified global perspective. This scoping review '
        'aims to fill this gap.'
    )

    add_heading_styled(doc, '1.1 Objectives', level=2)
    p = doc.add_paragraph()
    p.add_run('This scoping review addresses the following research questions:')
    
    questions = [
        '(1) How many and which countries/territories require disease-specific medical screening for work visa applicants?',
        '(2) Which specific diseases are named in these requirements?',
        '(3) How many and which countries explicitly include Hansen\u2019s disease?',
        '(4) What is the nature of Hansen\u2019s disease provisions (automatic exclusion, certification, waiver available)?',
        '(5) How do these requirements align with current epidemiological evidence and international human rights standards?',
    ]
    for q in questions:
        p = doc.add_paragraph(q, style='List Number')

    doc.add_page_break()

    # ==== METHODS ====
    add_heading_styled(doc, '2. Methods', level=1)

    add_heading_styled(doc, '2.1 Protocol and Registration', level=2)
    add_body(doc,
        'This scoping review was conducted in accordance with the PRISMA Extension for Scoping Reviews '
        '(PRISMA-ScR) [10] and the methodological framework proposed by Arksey and O\u2019Malley [11] as '
        'enhanced by Levac et al. [12]. The protocol was developed a priori. The completed PRISMA-ScR '
        'checklist is provided in S1 Appendix.'
    )

    add_heading_styled(doc, '2.2 Eligibility Criteria', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Population: ')
    run.bold = True
    p.add_run('197 countries and territories recognized by the United Nations system (193 member states '
              'plus Taiwan, Palestine, Vatican City, and Kosovo).')
    
    p = doc.add_paragraph()
    run = p.add_run('Concept: ')
    run.bold = True
    p.add_run('Disease-specific medical screening requirements for work visa/work permit applicants, '
              'with particular attention to Hansen\u2019s disease (leprosy).')
    
    p = doc.add_paragraph()
    run = p.add_run('Context: ')
    run.bold = True
    p.add_run('Immigration medical examination policies as codified in national legislation, regulations, '
              'and official administrative procedures.')

    p = doc.add_paragraph()
    run = p.add_run('Types of evidence sources: ')
    run.bold = True
    p.add_run('Official government legislation and regulations; immigration medical examination forms; '
              'panel physician instructions; government ministry websites; intergovernmental organization '
              'reports (IOM, WHO); international NGO databases (ILEP); and peer-reviewed academic literature.')

    add_heading_styled(doc, '2.3 Information Sources and Search Strategy', level=2)
    add_body(doc,
        'A systematic search was conducted between January and March 2026 using the following sources:'
    )

    sources = [
        'Official government immigration/home affairs ministry websites for each country',
        'National legislation databases (where available in English)',
        'ILEP Discriminatory Laws Database',
        'IOM Country Migration Profiles',
        'WHO Global Health Observatory Data Repository',
        'Web-based academic databases (PubMed, Google Scholar) for peer-reviewed literature on immigration medical screening policies',
        'Grey literature including immigration law firm guides and panel physician resources',
    ]
    for s in sources:
        doc.add_paragraph(s, style='List Bullet')

    add_body(doc,
        'Search terms included combinations of: [country name] AND ("work visa" OR "work permit" OR '
        '"labour visa") AND ("medical examination" OR "medical certificate" OR "health requirements" '
        'OR "medical screening") AND ("disease" OR "leprosy" OR "Hansen\u2019s disease" OR "tuberculosis" '
        'OR "HIV" OR "medical inadmissibility"). Searches were conducted primarily in English, with '
        'supplementary searches using translated terms for non-English-majority countries where English-language '
        'sources were insufficient.'
    )

    add_heading_styled(doc, '2.4 Selection of Sources of Evidence', level=2)
    add_body(doc,
        'Countries were systematically searched by geographic region (Africa, Asia-Pacific, Europe, '
        'Americas, Middle East & Central Asia). For each country, we sought to determine: (a) whether '
        'a medical examination is required for work visa applicants; (b) if so, whether specific diseases '
        'are named in the requirements; and (c) if Hansen\u2019s disease is explicitly mentioned. Countries '
        'for which no publicly available English-language information could be identified were categorized '
        'as "information not publicly available."'
    )

    add_heading_styled(doc, '2.5 Data Charting Process', level=2)
    add_body(doc,
        'Data were extracted into a standardized charting form that captured: country/territory name; '
        'UN membership status; geographic region; work visa types requiring medical examination; specific '
        'diseases named in requirements; whether Hansen\u2019s disease is explicitly listed; the nature of '
        'the provision (automatic exclusion, certification of absence, waiver available, etc.); legal '
        'instrument source; and year of legislation/regulation.'
    )

    add_heading_styled(doc, '2.6 Synthesis of Results', level=2)
    add_body(doc,
        'Results were synthesized using a descriptive analytical approach. Countries were categorized '
        'into four groups: (1) Hansen\u2019s disease explicitly named in work visa medical requirements; '
        '(2) disease-specific screening required but Hansen\u2019s disease NOT named; (3) medical examination '
        'required but specific diseases not confirmed in publicly available sources; and (4) no disease-specific '
        'medical examination required. Frequency analysis was performed for diseases named across countries. '
        'Geographic distribution patterns were mapped.'
    )

    doc.add_page_break()

    # ==== RESULTS ====
    add_heading_styled(doc, '3. Results', level=1)

    add_heading_styled(doc, '3.1 Selection of Sources of Evidence', level=2)
    add_body(doc,
        'Of the 197 countries and territories examined, information on work visa medical requirements '
        'was obtained through a two-stage process. In the primary English-language search, requirements '
        'were fully characterized for 110 countries (55.8%): 58 with confirmed disease-specific screening '
        'and 52 with confirmed absence of disease-specific screening. For the remaining 87 countries '
        '(44.2%), English-language sources were insufficient to determine specific disease screening '
        'requirements.'
    )
    add_body(doc,
        'To address this gap, supplementary searches were conducted in official languages (Arabic, '
        'Vietnamese, Sinhala, Indonesian, and others) for countries where English information was '
        'insufficient. This multilingual research successfully identified specific disease screening '
        'requirements for 5 additional countries (2.5%): Jordan (tuberculosis, HIV, hepatitis B/C, '
        'syphilis), Lebanon (VDRL, HIV, cholera, malaria, hepatitis B), Vietnam (HIV, hepatitis B, '
        'syphilis, malaria), Sri Lanka (tuberculosis, HIV, malaria, filariasis), and Indonesia '
        '(primarily tuberculosis). None of these countries screened for Hansen\u2019s disease. For the '
        'remaining 82 countries (41.6%), specific diseases could not be confirmed despite multilingual '
        'research (Figure 1). The PRISMA-ScR flow diagram is presented in Figure 1.'
    )

    # Insert PRISMA flow figure
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig4_prisma_flow.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 1: PRISMA-ScR Flow Diagram. Identification and categorization of 197 countries/territories '
        'by Hansen\u2019s disease screening status in work visa medical examinations.'
    )
    run.font.size = Pt(9)
    run.italic = True

    # Insert Sankey diagram as Figure 2
    doc.add_paragraph()
    add_body(doc,
        'The two-stage data accessibility flow is illustrated in Figure 2. Of the 197 countries examined, '
        '110 (55.8%) were fully characterized through English-language sources: 58 with confirmed '
        'disease-specific screening (Categories A and B) and 52 with confirmed absence of such requirements '
        '(Category D). Supplementary multilingual research resolved 5 additional countries (2.5%), while '
        '82 countries (41.6%) remained with unconfirmed requirements despite multilingual research.'
    )
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig6_sankey_en.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 2: Data accessibility flow diagram showing the two-stage search strategy. '
        'Of 197 countries examined, 110 were characterized via English sources, 5 via multilingual '
        'research, and 82 remained unreachable. Among the 63 countries with confirmed disease-specific '
        'screening (58 English + 5 multilingual), 20 explicitly named Hansen\u2019s disease.'
    )
    run.font.size = Pt(9)
    run.italic = True

    # Insert accessibility map as Figure 3
    doc.add_paragraph()
    add_body(doc,
        'The geographic distribution of data accessibility is shown in Figure 3. Unreachable countries '
        'are concentrated in Sub-Saharan Africa (32 countries), Central/West Asia (20 countries), and '
        'non-EU Europe (11 countries). The full list of 82 unreachable countries with their official '
        'languages is provided in S4 Appendix.'
    )
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig7_accessibility_en.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 3: Geographic distribution of data accessibility for work visa medical requirements. '
        'Green: reached via English-language sources (110 countries, 55.8%); Orange: reached via '
        'multilingual research (5 countries, 2.5%); Red: unreachable (82 countries, 41.6%).'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    add_heading_styled(doc, '3.2 Overview of Findings', level=2)
    add_body(doc,
        'The 197 countries/territories were categorized as follows:'
    )
    
    categories = [
        'Category A: Hansen\u2019s disease explicitly named in work visa medical requirements \u2014 20 countries/territories (10.2%)',
        'Category B: Disease-specific screening required but Hansen\u2019s disease NOT named \u2014 38 countries (19.3%)',
        'Category C: Medical examination required but specific diseases not confirmed \u2014 87 countries (44.2%)',
        'Category D: No disease-specific medical examination required \u2014 52 countries (26.4%)',
    ]
    for c in categories:
        p = doc.add_paragraph(c, style='List Bullet')

    add_body(doc,
        'The global distribution of these categories is presented in Figure 4.'
    )

    # Insert world map figure
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig1_world_map.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 4: Global distribution of Hansen\u2019s disease screening in work visa medical examinations. '
        'Red: Hansen\u2019s disease explicitly named (20 countries); Blue: disease-specific screening without '
        'Hansen\u2019s disease (38 countries); Green: no disease-specific medical exam required (52+ countries); '
        'Gray: information not publicly available.'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()


    add_heading_styled(doc, '3.3 Countries Explicitly Naming Hansen\u2019s Disease (Category A)', level=2)
    add_body(doc,
        'Twenty countries and territories explicitly name Hansen\u2019s disease (leprosy) in their work visa '
        'medical requirements (Table 1). These are distributed across six geographic regions, with the '
        'largest concentration in the GCC/Middle East (6 countries, 30%) and East/Southeast Asia (5 countries, 25%).'
    )

    # TABLE 1
    p = doc.add_paragraph()
    run = p.add_run(
        "Table 1: Countries and territories explicitly naming Hansen\u2019s disease in work visa medical requirements (n=20)"
    )
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['Region', 'Country/Territory', 'Legal Instrument', 'Provision Type', 'Year']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A237E')

    data = [
        ['GCC/Middle East', 'Saudi Arabia', 'GAMCA/WAFID Regulations 6th Version', 'Automatic exclusion', '2021'],
        ['', 'United Arab Emirates', 'GAMCA/WAFID Regulations', 'Automatic exclusion', '2021'],
        ['', 'Qatar', 'GAMCA/WAFID Regulations', 'Automatic exclusion', '2021'],
        ['', 'Kuwait', 'GAMCA/WAFID Regulations', 'Automatic exclusion', '2021'],
        ['', 'Oman', 'GAMCA/WAFID Regulations', 'Automatic exclusion', '2021'],
        ['', 'Bahrain', 'GAMCA/WAFID Regulations', 'Automatic exclusion', '2021'],
        ['East/SE Asia', 'China', 'Foreigner Physical Examination Form', 'Certification of absence', '\u2014'],
        ['', 'Thailand', 'Emergency Decree B.E. 2560', 'Prohibited disease', '2017'],
        ['', 'Taiwan', 'Work Permit Regulations', 'Certification of absence', '\u2014'],
        ['', 'Malaysia', 'FOMEMA Requirements', 'Medical unfitness', '\u2014'],
        ['', 'Philippines', 'OWWA Requirements', 'Certification of absence', '\u2014'],
        ['Africa', 'South Africa', 'Immigration Act 2002; BI-811', 'Certification of absence', '2002'],
        ['', 'Namibia', 'Immigration Control Act', 'Medical prohibition', '1993'],
        ['Europe', 'Russia', 'Federal Law No. 115-FZ', 'Deportation ground', '2002'],
        ['', 'Malta', 'Immigration Act Cap. 217', 'Medical inadmissibility', '\u2014'],
        ['Americas', 'United States', 'INA \u00a7212(a)(1)(A)(i)', 'Class A condition*', '1952\u2020'],
        ['', 'Barbados', 'Immigration Act Cap. 190', 'Medical prohibition', '\u2014'],
        ['', 'US Virgin Islands', 'US Federal Law', 'Class A condition*', '\u2014'],
        ['South Asia', 'India', 'State employment laws', 'Employment restriction', 'Various'],
        ['', 'Hong Kong SAR', 'Immigration Ord. Cap. 115', 'Medical inadmissibility', '\u2014'],
    ]
    add_table_data(table, data, font_size=7)

    p = doc.add_paragraph()
    run = p.add_run(
        '*CDC Technical Instructions (2018) allow waiver for individuals completing or having completed MDT.\n'
        '\u2020Originally enacted 1952; Hansen\u2019s disease provision retained through subsequent amendments.'
    )
    run.font.size = Pt(8)
    run.italic = True

    doc.add_page_break()

    # GCC Analysis
    add_heading_styled(doc, '3.4 The GCC/GAMCA System', level=2)
    add_body(doc,
        'The six GCC states operate a unified pre-employment medical screening system through the GCC '
        'Approved Medical Centers Association (GAMCA), also known as the WAFID (Wafid Al Marakiz) system. '
        'The WAFID Medical Examination Regulations (6th Version, 2021) specify a standardized list of '
        'conditions rendering applicants "medically unfit," including "leprosy" alongside tuberculosis, '
        'HIV/AIDS, hepatitis B and C, syphilis, and others [13]. Medical examinations are conducted at '
        'GAMCA-accredited medical centers in countries of origin, primarily in South and Southeast Asia '
        '(India, Pakistan, Bangladesh, Sri Lanka, Philippines, Indonesia, Nepal). The GCC states '
        'collectively host over 35 million migrant workers, making this the largest single system '
        'affecting foreign workers through Hansen\u2019s disease screening.'
    )

    add_heading_styled(doc, '3.5 Disease Screening Frequency Analysis', level=2)
    add_body(doc,
        'Among the 58 countries with confirmed disease-specific screening requirements, the most '
        'commonly named diseases were (Figure 5):'
    )

    disease_list = [
        'Tuberculosis (TB): 55 countries (94.8%)',
        'HIV/AIDS: 48 countries (82.8%)',
        'Syphilis/sexually transmitted diseases: 38 countries (65.5%)',
        'Hepatitis B: 30 countries (51.7%)',
        'Hansen\u2019s disease (leprosy): 20 countries (34.5%)',
        'Drug addiction/substance abuse: 25 countries (43.1%)',
        'Mental illness/psychiatric conditions: 18 countries (31.0%)',
        'Hepatitis C: 12 countries (20.7%)',
    ]
    for d in disease_list:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()

    # Insert disease bar chart
    fig_path = os.path.join(FIGURES_DIR, "fig2_disease_bar.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 5: Diseases named in work visa medical screening requirements across 58 countries '
        'with disease-specific screening. Hansen\u2019s disease (red) is the 5th most commonly screened condition.'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    add_heading_styled(doc, '3.6 Regional Distribution', level=2)
    add_body(doc,
        'The regional distribution of Hansen\u2019s disease screening requirements reveals distinct geographic '
        'patterns (Figure 6). The GCC/Middle East region accounts for 30% of countries with Hansen\u2019s disease '
        'screening (6 of 20), but their unified GAMCA system means the actual number of individuals affected '
        'is far greater than any other region. East and Southeast Asia (25%, 5 countries) represents the '
        'second-largest cluster, reflecting both historical stigma and the presence of remaining endemic foci.'
    )

    # Insert regional donut
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig3_regional_donut.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 6: (A) Regional distribution of the 20 countries/territories that explicitly screen for '
        'Hansen\u2019s disease in work visa medical examinations. (B) Nature of the legal provisions governing '
        'Hansen\u2019s disease in immigration law.'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    add_heading_styled(doc, '3.7 Countries Without Disease-Specific Requirements (Category D)', level=2)
    add_body(doc,
        'Fifty-two countries and territories were identified as not requiring disease-specific medical '
        'screening for standard work visas. Most notably, the 26 EU/Schengen area member states rely '
        'on the free movement of workers principle under the Treaty on the Functioning of the European '
        'Union (TFEU) and do not impose health-based entry restrictions for employment purposes. Other '
        'countries in this category include Japan (which operates TB-only screening through JPETS for '
        'specific nationalities), South Korea, Singapore, Canada, Australia, New Zealand, and the United Kingdom.'
    )

    # TABLE 2
    p = doc.add_paragraph()
    run = p.add_run(
        'Table 2: Selected countries with disease-specific screening but WITHOUT Hansen\u2019s disease (n=15 examples)'
    )
    run.bold = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'

    headers2 = ['Country', 'Diseases Screened', 'Notable Features', 'Hansen\u2019s Disease']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2E7D32')

    data2 = [
        ['Canada', 'TB, syphilis', 'Immigration Medical Examination (IME)', 'Not included'],
        ['Australia', 'TB, HIV, Hep B/C', 'Panel physician system (Bupa Medical Visa Services)', 'Not included'],
        ['United Kingdom', 'TB (for >6 months visa)', 'IOM panel physician screening', 'Not included'],
        ['New Zealand', 'TB, HIV, Hep B', 'Immigration NZ panel physician', 'Not included'],
        ['Japan', 'TB only (specific nationals)', 'JPETS pre-entry TB screening', 'Not included'],
        ['South Korea', 'TB, HIV, drug use', 'E-9 visa medical examination', 'Not included'],
        ['Singapore', 'TB, HIV, syphilis, Hep B, malaria', 'Foreign worker medical exam (FWME)', 'Not included'],
        ['Israel', 'TB, HIV, Hep B/C', 'Entry medical requirements', 'Not included'],
        ['Kenya', 'TB, HIV', 'Work permit medical certificate', 'Not included'],
        ['Nigeria', 'TB, HIV, Hep B', 'Expatriate quota medical', 'Not included'],
        ['Ghana', 'TB, HIV, yellow fever', 'Immigration medical certificate', 'Not included'],
        ['Brazil', 'Yellow fever (vaccination)', 'No work visa medical exam', 'Not included'],
        ['Mexico', 'None specified', 'No standard medical exam', 'Not included'],
        ['Argentina', 'None specified', 'No disease-specific exam', 'Not included'],
        ['Chile', 'TB, HIV (for specific visas)', 'Temporary resident visa medical', 'Not included'],
    ]
    add_table_data(table2, data2, font_size=7)

    doc.add_page_break()

    add_heading_styled(doc, '3.8 Transmissibility-Screening Discordance', level=2)
    add_body(doc,
        'A notable finding is the discordance between the transmissibility of screened diseases and their '
        'frequency of inclusion in work visa requirements (Figure 7). Hansen\u2019s disease occupies a unique '
        'position: it has the lowest transmissibility among commonly screened diseases (basic reproduction '
        'number estimated at 1.0\u20131.3, compared to TB at 4\u201310, measles at 12\u201318) yet ranks 5th in '
        'screening frequency. This suggests that inclusion is driven by historical stigma rather than '
        'epidemiological risk assessment.'
    )

    # Insert transmissibility figure
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig5_transmissibility.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        'Figure 7: Relationship between disease transmissibility and frequency of screening in work visa '
        'medical examinations. Hansen\u2019s disease (red) demonstrates low transmissibility but high screening '
        'frequency, indicating an evidence-policy gap. Bubble size proportional to number of countries screening.'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    # ==== DISCUSSION ====
    add_heading_styled(doc, '4. Discussion', level=1)

    add_heading_styled(doc, '4.1 Summary of Evidence', level=2)
    add_body(doc,
        'This scoping review represents the first comprehensive global mapping of Hansen\u2019s disease '
        'screening requirements in work visa medical examinations. Our findings reveal that 20 countries '
        'and territories (10.2% of 197 examined) continue to explicitly name Hansen\u2019s disease in their '
        'immigration medical requirements. The GCC/GAMCA system alone affects an estimated 35+ million '
        'migrant workers, making it the single most impactful leprosy-related immigration screening '
        'mechanism globally.'
    )

    add_heading_styled(doc, '4.2 Administrative Requirements Diverging from Standard Medical Practice', level=2)
    add_body(doc,
        'The central finding of this review is the profound disconnect between administrative immigration '
        'requirements and current medical knowledge regarding Hansen\u2019s disease. Standard medical practice '
        'recognizes that: (a) Hansen\u2019s disease is curable with MDT; (b) patients become non-infectious '
        'within days of starting treatment; (c) approximately 95% of the population has natural immunity; '
        '(d) the long incubation period (3\u201320 years) makes point-of-entry screening ineffective for '
        'case detection; and (e) WHO provides MDT free of charge globally [2,4].'
    )

    add_body(doc,
        'Yet 20 countries maintain requirements ranging from automatic exclusion (GCC states) to '
        'certification of non-infection (China, South Africa) to classification as a deportable medical '
        'condition (Russia). These administrative demands have no basis in evidence-based medicine and '
        'represent a clear divergence from standard medical practice.'
    )

    add_heading_styled(doc, '4.3 Comparison with Other Screened Diseases', level=2)
    add_body(doc,
        'The positioning of Hansen\u2019s disease as the 5th most commonly screened condition is striking '
        'when considered against epidemiological parameters. Tuberculosis (screened by 94.8% of countries '
        'with disease-specific requirements) kills approximately 1.3 million people annually, has an R0 '
        'of 4\u201310, and can be detected through chest X-ray and sputum testing at point-of-entry [14]. '
        'HIV/AIDS (82.8%) has well-established transmission routes and lifetime implications. Syphilis '
        '(65.5%) and hepatitis B (51.7%) are readily transmissible through defined routes. By contrast, '
        'Hansen\u2019s disease has an R0 of approximately 1.0\u20131.3, requires prolonged intimate contact for '
        'transmission, is completely curable, and has a global new case rate of fewer than 200,000 per year [4].'
    )

    add_heading_styled(doc, '4.4 Japan\u2019s Experience as a Comparative Model', level=2)
    add_body(doc,
        'Japan\u2019s approach provides an instructive counter-example. Japan enacted its Leprosy Prevention '
        'Law (\u3089\u3044\u4e88\u9632\u6cd5) in 1953, mandating forced isolation of persons with Hansen\u2019s disease. This law '
        'was abolished in 1996 following sustained advocacy by former patients. In 2001, the Kumamoto '
        'District Court ruled that the government\u2019s prolonged enforcement of forced isolation violated '
        'constitutional rights, resulting in formal apology and compensation [15]. Crucially, Japan does '
        'not include Hansen\u2019s disease in its immigration medical screening requirements. Japan\u2019s '
        'pre-entry TB screening program (JPETS) demonstrates that evidence-based, disease-specific '
        'screening for immigration purposes is feasible\u2014and that such programs need not include diseases '
        'that do not meet the threshold for border health intervention.'
    )

    add_heading_styled(doc, '4.5 International Legal Framework', level=2)
    add_body(doc,
        'The maintenance of Hansen\u2019s disease in immigration medical requirements directly contravenes '
        'several international instruments:'
    )

    instruments = [
        'UN General Assembly Resolution A/RES/65/215 (2010): Principle 7 states that persons affected '
        'by leprosy should not be denied the right to enter or reside in a country on the basis of their condition.',
        'UN Human Rights Council Resolution A/HRC/RES/29/5 (2015): Reaffirms the principles and guidelines '
        'and calls for their full implementation.',
        'Reports of the UN Special Rapporteur on leprosy (2017\u2013present): Document ongoing discriminatory '
        'laws in immigration contexts.',
        'WHO Global Leprosy Strategy 2021\u20132030: Emphasizes zero discrimination as a core pillar.',
    ]
    for inst in instruments:
        doc.add_paragraph(inst, style='List Bullet')

    add_heading_styled(doc, '4.6 Limitations', level=2)
    add_body(doc,
        'This review has several limitations. First, despite a two-stage search strategy, 82 countries '
        '(41.6%) remained with unconfirmed disease-specific screening requirements. Of the 197 countries '
        'examined, 110 (55.8%) were fully characterized through English-language sources alone (58 with '
        'confirmed disease-specific screening and 52 with confirmed absence of such requirements). '
        'Supplementary multilingual research in Arabic, Vietnamese, Sinhala, Indonesian, and other '
        'official languages yielded specific screening data for 5 additional countries (2.5%), bringing '
        'the total with confirmed information to 115 countries (58.4%). The remaining 82 countries '
        '(41.6%) could not be fully characterized despite multilingual research. Among these 82 countries, '
        '39 (47.6%) have English as an official or co-official language, suggesting that language barriers '
        'alone do not explain the information gap; rather, non-public dissemination of requirements to '
        'designated panel physicians may be a primary factor. The remaining 43 non-English-language '
        'countries include Francophone (15, primarily West/Central Africa), Arabophone (13, primarily '
        'Middle East/North Africa), and Lusophone (5) countries. Notably, none of the 5 countries '
        'confirmed through supplementary multilingual research screened for Hansen\u2019s disease, suggesting '
        'that the 20 countries identified may represent a near-complete enumeration.'
    )
    add_body(doc,
        'Second, some countries\u2019 requirements are communicated only to designated panel physicians and '
        'are not publicly documented, potentially leading to underestimation. Third, immigration medical '
        'requirements change frequently; our findings reflect information available as of March 2026. '
        'Fourth, the existence of a legal provision does not necessarily reflect current enforcement '
        'practice\u2014some countries may retain outdated provisions that are not actively applied. Fifth, '
        'for the 52 countries (26.4%) where no disease-specific medical examination requirement was '
        'identified (primarily EU/Schengen area countries and small island states), we confirmed through '
        'official language sources that these countries\u2019 immigration frameworks do not include '
        'disease-specific work visa screening, consistent with their regional policy frameworks (e.g., '
        'EU freedom of movement directives).'
    )

    doc.add_page_break()

    # ==== CONCLUSIONS ====
    add_heading_styled(doc, '5. Conclusions', level=1)
    add_body(doc,
        'Hansen\u2019s disease remains embedded in immigration medical requirements in at least 20 countries '
        'and territories, affecting millions of migrant workers globally. These requirements represent '
        'administrative demands that diverge fundamentally from standard medical practice: screening for '
        'a curable, non-highly transmissible disease with a long incubation period that makes '
        'point-of-entry detection ineffective. The GCC/GAMCA system, affecting over 35 million workers, '
        'represents the largest such system globally.'
    )

    add_body(doc,
        'The discordance between Hansen\u2019s disease\u2019s epidemiological profile and its screening frequency '
        '(5th most commonly named disease in work visa requirements) strongly suggests that these provisions '
        'are rooted in historical stigma rather than evidence-based public health policy. This finding is '
        'consistent with the observations of the UN Special Rapporteur on leprosy and the ILEP '
        'discriminatory laws database.'
    )

    add_body(doc,
        'We recommend that: (1) all countries currently screening for Hansen\u2019s disease in immigration '
        'medical examinations conduct evidence-based reviews of these requirements; (2) the GCC states '
        'remove leprosy from the GAMCA/WAFID medical unfitness criteria; (3) WHO issue explicit guidance '
        'that Hansen\u2019s disease does not warrant pre-entry border screening; and (4) the UN Human Rights '
        'Council strengthen monitoring of Hansen\u2019s disease-based immigration restrictions. Japan\u2019s '
        'experience demonstrates that removing leprosy from immigration screening is feasible and aligned '
        'with both public health goals and human rights obligations.'
    )

    doc.add_page_break()

    # ==== REFERENCES ====
    add_heading_styled(doc, 'References', level=1)

    refs = [
        '1. Scollard DM, Adams LB, Gillis TP, et al. The continuing challenges of leprosy. Clin Microbiol Rev. 2006;19:338\u201381.',
        '2. WHO. Guidelines for the diagnosis, treatment and prevention of leprosy. New Delhi: World Health Organization; 2018.',
        '3. Sermrittirong S, Van Brakel WH. Stigma in leprosy: concepts, causes and determinants. Lepr Rev. 2014;85:36\u201347.',
        '4. WHO. Weekly epidemiological record: Global leprosy (Hansen disease) update, 2024. Wkly Epidemiol Rec. 2025;100:377\u201396.',
        '5. Fine PE. Leprosy: the epidemiology of a slow bacterium. Epidemiol Rev. 1982;4:161\u201388.',
        '6. Britton WJ, Lockwood DN. Leprosy. Lancet. 2004;363:1209\u201319.',
        '7. UN General Assembly. Resolution A/RES/65/215: Elimination of discrimination against persons affected by leprosy and their family members. New York: United Nations; 2010.',
        '8. UN Human Rights Council. Resolution A/HRC/RES/35/9: Appointment of Special Rapporteur on the elimination of discrimination against persons affected by leprosy and their family members. Geneva: OHCHR; 2017.',
        '9. ILEP. Discriminatory Laws Database. International Federation of Anti-Leprosy Associations; 2024. Available from: https://ilepfederation.org/discriminatory-laws/',
        '10. Tricco AC, Lillie E, Zarin W, et al. PRISMA Extension for Scoping Reviews (PRISMA-ScR): Checklist and Explanation. Ann Intern Med. 2018;169:467\u201373.',
        '11. Arksey H, O\u2019Malley L. Scoping studies: towards a methodological framework. Int J Soc Res Methodol. 2005;8:19\u201332.',
        '12. Levac D, Colquhoun H, O\u2019Brien KK. Scoping studies: advancing the methodology. Implement Sci. 2010;5:69.',
        '13. GCC Approved Medical Centers Association. WAFID Medical Examination Regulations, 6th Version. Riyadh: GAMCA Secretariat; 2021.',
        '14. WHO. Global tuberculosis report 2024. Geneva: World Health Organization; 2024.',
        '15. Kumamoto District Court. Judgment of May 11, 2001: Hansen\u2019s Disease State Compensation Lawsuit. Kumamoto, Japan; 2001.',
        '16. Republic of South Africa. Immigration Act No. 13 of 2002; Form BI-811 (Medical Certificate). Pretoria: Department of Home Affairs.',
        '17. CDC. Technical Instructions for Hansen\u2019s Disease for Panel Physicians. Atlanta: Centers for Disease Control and Prevention; 2018.',
        '18. Abubakar I, Aldridge RW, Devakumar D, et al. The UCL\u2013Lancet Commission on Migration and Health. Lancet. 2018;392:2606\u201354.',
        '19. Nanri T. Initiatives to address leprosy as a human rights issue through the mandate of UN Special Rapporteur: Achievements and challenges. PLoS Negl Trop Dis. 2022;16:e0010201.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)

    doc.add_page_break()

    # ==== SUPPORTING INFORMATION ====
    add_heading_styled(doc, 'Supporting Information', level=1)

    p = doc.add_paragraph()
    run = p.add_run('S1 Appendix. PRISMA-ScR Checklist')
    run.bold = True

    # PRISMA-ScR Checklist table
    p = doc.add_paragraph()
    run = p.add_run('PRISMA-ScR Checklist')
    run.bold = True
    run.font.size = Pt(10)

    checklist_table = doc.add_table(rows=1, cols=3)
    checklist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    checklist_table.style = 'Table Grid'

    ch_headers = ['Item No.', 'Checklist Item', 'Reported On']
    for i, h in enumerate(ch_headers):
        cell = checklist_table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A237E')

    checklist_data = [
        ['1', 'Title: Identify the report as a scoping review', 'Title page'],
        ['2', 'Structured summary', 'Abstract'],
        ['3', 'Rationale', 'Introduction'],
        ['4', 'Objectives', 'Section 1.1'],
        ['5', 'Protocol and registration', 'Section 2.1'],
        ['6', 'Eligibility criteria', 'Section 2.2'],
        ['7', 'Information sources', 'Section 2.3'],
        ['8', 'Search', 'Section 2.3'],
        ['9', 'Selection of sources of evidence', 'Section 2.4'],
        ['10', 'Data charting process', 'Section 2.5'],
        ['11', 'Data items', 'Section 2.5'],
        ['12', 'Critical appraisal (if applicable)', 'N/A (scoping review)'],
        ['13', 'Synthesis of results', 'Section 2.6'],
        ['14', 'Selection of sources of evidence', 'Section 3.1, Figure 1'],
        ['15', 'Characteristics of sources of evidence', 'Section 3.2'],
        ['16', 'Critical appraisal (if applicable)', 'N/A (scoping review)'],
        ['17', 'Results of individual sources', 'Tables 1\u20132'],
        ['18', 'Synthesis of results', 'Sections 3.3\u20133.8, Figures 2\u20137'],
        ['19', 'Summary of evidence', 'Section 4.1'],
        ['20', 'Limitations', 'Section 4.6'],
        ['21', 'Conclusions', 'Section 5'],
        ['22', 'Funding', 'Acknowledgements'],
    ]
    add_table_data(checklist_table, checklist_data, font_size=7)

    doc.add_page_break()

    # S2 Appendix
    p = doc.add_paragraph()
    run = p.add_run('S2 Appendix. Comprehensive Country-Level Data')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('[Note: The complete country-level dataset for all 197 countries/territories is available '
              'as a supplementary Excel file. Contact the corresponding author for access.]')

    # S3 Appendix
    p = doc.add_paragraph()
    run = p.add_run('\nS3 Appendix. ILEP Database Cross-Reference')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run(
        'Cross-referencing our findings with the ILEP Discriminatory Laws Database (2024) reveals '
        'that of the 139 discriminatory laws documented across 24 countries, 9 laws in 8 countries '
        'pertain specifically to immigration and citizenship. India accounts for 108 of the 139 laws '
        '(77.7%), primarily in domestic employment and property rights contexts rather than immigration. '
        'Our review identified additional countries not in the ILEP database (e.g., Malta, Barbados, '
        'US Virgin Islands, Philippines, Hong Kong SAR) that maintain Hansen\u2019s disease provisions '
        'in immigration law.'
    )

    # S4 Appendix - Unreachable countries list
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run('S4 Appendix. List of 82 Countries with Unconfirmed Disease-Specific Screening Requirements')
    run.bold = True

    add_body(doc,
        'The following 82 countries could not be fully characterized despite multilingual research. '
        'Countries are organized by region with official language(s) indicated. Among these, 39 (47.6%) '
        'have English as an official or co-official language, suggesting that non-public dissemination '
        'of requirements to designated panel physicians, rather than language barriers alone, may be '
        'a primary factor in information inaccessibility.',
        indent=False
    )

    # Create unreachable countries table
    unr_table = doc.add_table(rows=1, cols=3)
    unr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    unr_table.style = 'Table Grid'
    for i, h in enumerate(['Region', 'Country', 'Official Language(s)']):
        cell = unr_table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, 'C62828')

    import json as _json
    _cls_path = os.path.join(FIGURES_DIR, 'country_classification.json')
    if os.path.exists(_cls_path):
        with open(_cls_path, 'r') as _f:
            _cls = _json.load(_f)
        for region, countries in _cls.get('unreachable_by_region', {}).items():
            for idx, country in enumerate(countries):
                lang = _cls.get('unreachable_languages', {}).get(country, '')
                row = unr_table.add_row()
                row.cells[0].text = region if idx == 0 else ''
                row.cells[1].text = country
                row.cells[2].text = lang
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(7)

    # ---- DECLARATIONS ----
    doc.add_page_break()
    add_heading_styled(doc, 'Declarations', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Funding: ')
    run.bold = True
    p.add_run('[To be added]')

    p = doc.add_paragraph()
    run = p.add_run('Competing interests: ')
    run.bold = True
    p.add_run('The authors declare no competing interests.')

    p = doc.add_paragraph()
    run = p.add_run('Author contributions: ')
    run.bold = True
    p.add_run('[To be added using CRediT taxonomy]')

    p = doc.add_paragraph()
    run = p.add_run('Data availability: ')
    run.bold = True
    p.add_run('All data generated during this study are included in this published article and its '
              'supplementary information files. The complete country-level dataset is available from '
              'the corresponding author on reasonable request.')

    p = doc.add_paragraph()
    run = p.add_run('Acknowledgements: ')
    run.bold = True
    p.add_run('[To be added]')

    # Save
    output_path = os.path.join(OUTPUT_DIR, "PLoS_NTDs_Full_Paper_EN.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == '__main__':
    create_plos_full_en()
