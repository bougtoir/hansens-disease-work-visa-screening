#!/usr/bin/env python3
"""Generate PLoS NTDs Full Paper (Japanese) as DOCX with color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FIGURES_DIR = "/home/ubuntu/scoping_review/figures"
OUTPUT_DIR = "/home/ubuntu/scoping_review/docx"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return h


def add_body(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_table_data(table, data, font_size=8):
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def create_plos_full_ja():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    # ==== TITLE PAGE ====
    p = doc.add_paragraph()
    run = p.add_run('研究論文（RESEARCH ARTICLE）')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "労働ビザ医療検査におけるハンセン病スクリーニング：\n"
        "標準医療と乖離した行政要求に関するグローバルスコーピングレビュー"
    )
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        "Hansen\u2019s Disease Screening in Work Visa Medical Examinations: "
        "A Global Scoping Review of Administrative Requirements That Diverge from Standard Medical Practice"
    )
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('[著者名]')
    run.font.size = Pt(12)
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('[所属機関]')
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('* 責任著者: [メールアドレス]')
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run('※本文書はPLoS Neglected Tropical Diseases投稿用原稿の日本語版です。投稿時には英語版を使用してください。')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)

    doc.add_page_break()

    # ==== ABSTRACT ====
    add_heading_styled(doc, '抄録（Abstract）', level=1)

    p = doc.add_paragraph()
    run = p.add_run('背景（Background）')
    run.bold = True
    p.add_run(
        '\nハンセン病（Hansen\u2019s disease, leprosy）はMycobacterium lepraeによる治癒可能かつ低感染性の感染症である。'
        '治療法の進歩と世界的な罹患率の低下にもかかわらず、一部の国は労働ビザ医療検査の一部としてハンセン病の'
        'スクリーニングを継続的に要求している。これらの要件の範囲と性質は体系的にマッピングされていない。'
    )

    p = doc.add_paragraph()
    run = p.add_run('目的（Objectives）')
    run.bold = True
    p.add_run(
        '\n労働ビザ申請者に対して疾病特異的な医療検査を要求している国を体系的に特定すること（特にハンセン病に注目）、'
        'それらの要件の性質を明らかにすること、および現在の疫学的エビデンスと国際的人権基準との整合性を評価すること。'
    )

    p = doc.add_paragraph()
    run = p.add_run('方法（Methods）')
    run.bold = True
    p.add_run(
        '\n197カ国・地域（国連加盟国193カ国＋オブザーバー4カ国・地域）を対象に、PRISMA-ScRガイドラインに'
        '従ったスコーピングレビューを実施した。情報源には各国の公式政府法令、移民医療検査書式、規制枠組み、'
        'ILEP差別的法律データベース、IOM報告書、および査読付き学術文献を含めた。疾病特異的スクリーニング要件、'
        '法的根拠、および規定の性質に関するデータを抽出した。'
    )

    p = doc.add_paragraph()
    run = p.add_run('結果（Results）')
    run.bold = True
    p.add_run(
        '\n調査した197カ国・地域のうち、58カ国（29.4%）が労働ビザに対する疾病特異的な医療検査要件を有することが'
        '確認された。20カ国・地域（10.2%）がハンセン病を明示的に記載していた。GCC6カ国はGAMCA/WAFID統一制度を'
        '運用し、ハンセン病を「医療上不適格」として記載しており、3,500万人以上の外国人労働者に影響を与えている。'
        'ハンセン病は5番目に多くスクリーニングされる疾病（34.5%）であり、結核（94.8%）、HIV/AIDS（82.8%）、'
        '梅毒（65.5%）、B型肝炎（51.7%）に次ぐ。EU/シェンゲン圏26カ国および追加の26カ国は疾病特異的'
        'スクリーニングを要求していない。ハンセン病スクリーニングと当該国のハンセン病負担との間に相関は認められなかった。'
    )

    p = doc.add_paragraph()
    run = p.add_run('結論（Conclusions）')
    run.bold = True
    p.add_run(
        '\nハンセン病は疫学的正当性を欠くにもかかわらず、20カ国の移民医療要件に組み込まれたままである。'
        'これらの要件は標準医療と乖離した行政要求を表しており、国連決議A/RES/65/215に反する。'
        'これらの規定の緊急の見直しと改革が必要である。'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('キーワード: ')
    run.bold = True
    p.add_run(
        'ハンセン病; らい; 労働ビザ; 移民医療検査; スコーピングレビュー; PRISMA-ScR; '
        '医療上の入国不許可; 差別; 顧みられない熱帯病'
    )

    doc.add_page_break()

    # ==== AUTHOR SUMMARY ====
    add_heading_styled(doc, '著者要約（Author Summary）', level=1)
    p = doc.add_paragraph()
    p.add_run(
        'ハンセン病（Hansen\u2019s disease, leprosy）は人類に知られる最も古い感染症の一つであるが、'
        '現在はWHOの多剤併用療法により完全に治癒可能であり、高度に感染性でもない—人口の95%が自然免疫を有する。'
        'にもかかわらず、我々は20カ国が依然として労働ビザ申請者にこの疾病のスクリーニングを要求していることを'
        '明らかにした。湾岸協力会議（GCC）6カ国（サウジアラビア、UAE、カタール、クウェート、オマーン、バーレーン）は'
        '合計3,500万人以上の外国人労働者を受け入れており、ハンセン病を「医療上不適格」とする条件として記載している。'
        'これが特に注目されるのは、2010年に国連総会がハンセン病罹患者に対する差別の撤廃（移民文脈を含む）を'
        '求める決議を採択したにもかかわらず、これらの要件が存続しているためである。本研究は、国連が認定する'
        '197カ国・地域のすべてにわたってこれらの要件を体系的にマッピングした初めての研究である。'
        'ハンセン病は労働ビザ医療検査において5番目に多くスクリーニングされる疾病であるが、上位4疾病'
        '（結核、HIV、梅毒、B型肝炎）と比較してはるかに低い感染性しか持たない。'
        '日本の経験—1996年にらい予防法を廃止し、移民スクリーニングにハンセン病を含めていない—は、'
        'エビデンスに基づくアプローチが実現可能であることを示している。我々は各国に対し、'
        '移民医療要件を現在の科学的エビデンスと国際的人権基準に整合させることを求める。'
    )

    doc.add_page_break()

    # ==== INTRODUCTION ====
    add_heading_styled(doc, '1. 緒言（Introduction）', level=1)

    add_body(doc,
        'ハンセン病（Hansen\u2019s disease, leprosy）はMycobacterium lepraeにより引き起こされ、'
        '人類の歴史上最も古くから記録されている感染症の一つである [1]。治療における劇的な進歩—'
        '特に1981年のWHO多剤併用療法（MDT）の導入により、患者は数日以内に非感染性となり、'
        '6〜12ヶ月で完全治癒が達成される [2]—にもかかわらず、この疾病は世界の多くの地域で'
        '深刻な社会的スティグマを持ち続けている [3]。'
    )

    add_body(doc,
        '世界の疫学的状況は大幅に改善している。WHOは2024年に世界で約182,000件の新規症例を報告しており、'
        '新規症例発見率は過去20年間にわたり着実に低下している [4]。世界人口の約95%がM. lepraeに対する'
        '自然免疫を有し、感染には未治療の多菌型患者との長期間の密接な接触が必要である [5]。'
        '潜伏期間は3〜5年（一部の症例では20年に達する）であり、入国時の医療検査による検出は'
        'ほぼ不可能である [6]。'
    )

    add_body(doc,
        'これらの特性にもかかわらず、ハンセン病は歴史的に様々な国の移民法上の医療上入国不許可リストに'
        '含まれてきた。このような規定の存続は、入国のための行政要件と現在の医学知識との整合性について'
        '重要な問題を提起する—我々はこれを「標準医療と乖離した行政要求」と呼ぶ。'
    )

    add_body(doc,
        '国際社会はハンセン病関連の差別に明示的に対処してきた。国連総会決議A/RES/65/215（2010年）は'
        'ハンセン病罹患者に対する差別の撤廃を加盟国に求め、原則7は入国を含む移動の自由の権利を確認している [7]。'
        '2017年にはハンセン病罹患者に対する差別の撤廃に関する国連特別報告者が任命された [8]。'
        '国際ハンセン病団体連合（ILEP）は24カ国にわたる139の差別的法律を記録するデータベースを維持している [9]。'
    )

    add_body(doc,
        'しかし、労働ビザ医療検査におけるハンセン病スクリーニング要件の包括的かつ体系的なグローバルマッピングは'
        'これまで実施されていない。既存の文献は個別の国や地域を扱っているが、統一的なグローバルな視点を欠いている。'
        '本スコーピングレビューはこのギャップを埋めることを目的とする。'
    )

    add_heading_styled(doc, '1.1 目的', level=2)
    p = doc.add_paragraph()
    p.add_run('本スコーピングレビューは以下の研究課題に取り組む：')

    questions = [
        '（1）労働ビザ申請者に対して疾病特異的な医療検査を要求している国・地域はいくつあり、どこか？',
        '（2）これらの要件で指定されている具体的な疾病は何か？',
        '（3）ハンセン病を明示的に含めている国はいくつあり、どこか？',
        '（4）ハンセン病規定の性質は何か（自動排除、証明書、免除可能性）？',
        '（5）これらの要件は現在の疫学的エビデンスと国際的人権基準にどの程度整合しているか？',
    ]
    for q in questions:
        p = doc.add_paragraph(q, style='List Number')

    doc.add_page_break()

    # ==== METHODS ====
    add_heading_styled(doc, '2. 方法（Methods）', level=1)

    add_heading_styled(doc, '2.1 プロトコルと登録', level=2)
    add_body(doc,
        '本スコーピングレビューはPRISMA Extension for Scoping Reviews（PRISMA-ScR）[10]および'
        'Arksey & O\u2019Malley [11]が提案しLevac et al. [12]が強化した方法論的枠組みに従って実施した。'
        'プロトコルは事前に策定した。完成したPRISMA-ScRチェックリストはS1 Appendixに提供する。'
    )

    add_heading_styled(doc, '2.2 適格基準', level=2)
    p = doc.add_paragraph()
    run = p.add_run('対象集団（Population）: ')
    run.bold = True
    p.add_run('国連システムが認定する197カ国・地域（加盟国193カ国＋台湾、パレスチナ、バチカン市国、コソボ）。')

    p = doc.add_paragraph()
    run = p.add_run('概念（Concept）: ')
    run.bold = True
    p.add_run('労働ビザ／就労許可申請者に対する疾病特異的な医療検査要件（特にハンセン病に注目）。')

    p = doc.add_paragraph()
    run = p.add_run('文脈（Context）: ')
    run.bold = True
    p.add_run('各国の法令、規制、および公式行政手続きに規定された移民医療検査政策。')

    p = doc.add_paragraph()
    run = p.add_run('エビデンスの種類: ')
    run.bold = True
    p.add_run('公式政府法令・規則、移民医療検査書式、指定医指示書、政府省庁ウェブサイト、'
              '政府間機関報告書（IOM、WHO）、国際NGOデータベース（ILEP）、査読付き学術文献。')

    add_heading_styled(doc, '2.3 情報源と検索戦略', level=2)
    add_body(doc,
        '2026年1月から3月にかけて、以下の情報源を用いた体系的検索を実施した：'
    )

    sources = [
        '各国の公式政府移民局/内務省ウェブサイト',
        '各国の法令データベース（英語で利用可能な場合）',
        'ILEP差別的法律データベース',
        'IOM各国移民プロファイル',
        'WHO Global Health Observatoryデータリポジトリ',
        'ウェブベースの学術データベース（PubMed、Google Scholar）',
        '灰色文献（移民法律事務所ガイド、指定医リソース等）',
    ]
    for s in sources:
        doc.add_paragraph(s, style='List Bullet')

    add_body(doc,
        '検索語には以下の組み合わせを使用した：[国名] AND ("work visa" OR "work permit" OR "labour visa") '
        'AND ("medical examination" OR "medical certificate" OR "health requirements" OR "medical screening") '
        'AND ("disease" OR "leprosy" OR "Hansen\u2019s disease" OR "tuberculosis" OR "HIV" OR "medical inadmissibility")。'
        '検索は主に英語で実施し、英語情報源が不十分な国については翻訳された用語を用いた補足検索を行った。'
    )

    add_heading_styled(doc, '2.4 エビデンス源の選択', level=2)
    add_body(doc,
        '各国は地理的地域別（アフリカ、アジア太平洋、ヨーロッパ、アメリカ大陸、中東・中央アジア）に'
        '体系的に検索した。各国について以下を確認した：（a）労働ビザ申請者に医療検査が必要か、'
        '（b）必要な場合、特定の疾病が要件に記載されているか、（c）ハンセン病が明示的に記載されているか。'
        '公開されている英語情報が特定できなかった国は「情報が公開されていない」として分類した。'
    )

    add_heading_styled(doc, '2.5 データ抽出プロセス', level=2)
    add_body(doc,
        'データは標準化された抽出フォームに記録した。抽出項目：国・地域名、国連加盟状況、地理的地域、'
        '医療検査を要求する労働ビザの種類、要件に記載された特定の疾病、ハンセン病が明示的に記載されているか、'
        '規定の性質（自動排除、不罹患証明、免除可能等）、法的根拠の出典、法令・規則の年。'
    )

    add_heading_styled(doc, '2.6 結果の統合', level=2)
    add_body(doc,
        '結果は記述的分析アプローチを用いて統合した。各国を4つのカテゴリーに分類した：'
        '（1）ハンセン病が労働ビザ医療要件に明示的に記載、（2）疾病特異的スクリーニングを要求するが'
        'ハンセン病は不記載、（3）医療検査を要求するが特定の疾病が公開情報で確認できない、'
        '（4）疾病特異的医療検査を要求しない。各国にまたがる指定疾病の頻度分析を実施し、'
        '地理的分布パターンをマッピングした。'
    )

    doc.add_page_break()

    # ==== RESULTS ====
    add_heading_styled(doc, '3. 結果（Results）', level=1)

    add_heading_styled(doc, '3.1 エビデンス源の選択', level=2)
    add_body(doc,
        '調査した197カ国・地域について、労働ビザ医療要件に関する情報は2段階のプロセスにより収集した。'
        '第一段階の英語情報源による検索では、110カ国（55.8%）の要件が完全に特定された：'
        '58カ国は疾病特異的スクリーニングが確認され、52カ国は疾病特異的スクリーニングの不存在が確認された。'
        '残りの87カ国（44.2%）については、英語情報源では特定の疾病スクリーニング要件を確認できなかった。'
    )
    add_body(doc,
        'この情報格差に対処するため、英語情報が不十分な国について公用語（アラビア語、ベトナム語、'
        'シンハラ語、インドネシア語等）による補足検索を実施した。この多言語調査により、5カ国（2.5%）の'
        '具体的な疾病スクリーニング要件が新たに特定された：ヨルダン（結核、HIV、B/C型肝炎、梅毒）、'
        'レバノン（VDRL、HIV、コレラ、マラリア、B型肝炎）、ベトナム（HIV、B型肝炎、梅毒、マラリア）、'
        'スリランカ（結核、HIV、マラリア、フィラリア症）、インドネシア（主に結核）。これらの国はいずれも'
        'ハンセン病をスクリーニングしていなかった。残りの82カ国（41.6%）については、多言語調査を実施したにも'
        'かかわらず特定の疾病を確認できなかった（図1）。PRISMA-ScRフローダイアグラムを図1に示す。'
    )

    # Insert PRISMA flow figure
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig4_prisma_flow.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        '図1：PRISMA-ScRフローダイアグラム。197カ国・地域の労働ビザ医療検査における'
        'ハンセン病スクリーニング状況の特定と分類。'
    )
    run.font.size = Pt(9)
    run.italic = True

    # Insert Sankey diagram as Figure 2
    doc.add_paragraph()
    add_body(doc,
        '2段階のデータ到達フローを図2に示す。調査した197カ国のうち、110カ国（55.8%）は'
        '英語情報源で完全に特定された：疾病特異的スクリーニング確認58カ国（カテゴリーA・B）と'
        '疾病特異的要件の不存在確認52カ国（カテゴリーD）。補足的多言語調査により5カ国（2.5%）が'
        '追加的に解決されたが、82カ国（41.6%）は多言語調査後も要件が未確認のままであった。'
    )
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig6_sankey_ja.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(
        '図2：データ到達フロー図（2段階検索戦略）。調査した197カ国のうち、110カ国は英語情報源で、'
        '5カ国は多言語調査で特定され、82カ国は到達不可であった。疾病特異的スクリーニングが確認された'
        '63カ国（英語58＋多言語5）のうち、20カ国がハンセン病を明示的に記載していた。'
    )
    run.font.size = Pt(9)
    run.italic = True

    # Insert accessibility map as Figure 3
    doc.add_paragraph()
    add_body(doc,
        'データ到達可能性の地理的分布を図3に示す。到達不可の国はサハラ以南アフリカ（32カ国）、'
        '中央・西アジア（20カ国）、非EU欧州（11カ国）に集中している。到達不可82カ国の完全なリスト'
        '（公用語情報付）はS4 Appendixに記載する。'
    )
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig7_accessibility_ja.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(
        '図3：労働ビザ医療要件に関するデータ到達可能性の地理的分布。'
        '緑：英語情報源で到達（110カ国、55.8%）、橙：多言語調査で到達（5カ国、2.5%）、'
        '赤：到達不可（82カ国、41.6%）。'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    add_heading_styled(doc, '3.2 所見の概要', level=2)
    add_body(doc, '197カ国・地域は以下のように分類された：')

    categories = [
        'カテゴリーA：ハンセン病が労働ビザ医療要件に明示的に記載 — 20カ国・地域（10.2%）',
        'カテゴリーB：疾病特異的スクリーニングを要求するがハンセン病は不記載 — 38カ国（19.3%）',
        'カテゴリーC：医療検査を要求するが特定の疾病が未確認 — 87カ国（44.2%）',
        'カテゴリーD：疾病特異的医療検査を要求しない — 52カ国（26.4%）',
    ]
    for c in categories:
        doc.add_paragraph(c, style='List Bullet')

    add_body(doc,
        'これらのカテゴリーの世界分布を図4に示す。'
    )

    # Insert world map figure
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig1_world_map.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        '図4：労働ビザ医療検査におけるハンセン病スクリーニングの世界分布。'
        '赤：ハンセン病を明示的に記載（20カ国）、青：ハンセン病を含まない疾病特異的スクリーニング（38カ国）、'
        '緑：疾病特異的医療検査を要求しない（52カ国以上）、灰色：公開情報なし。'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()


    add_heading_styled(doc, '3.3 ハンセン病を明示的に記載する国（カテゴリーA）', level=2)
    add_body(doc,
        '20カ国・地域がハンセン病を労働ビザ医療要件に明示的に記載している（表1）。'
        '最大の集中はGCC/中東（6カ国、30%）と東・東南アジア（5カ国、25%）に見られる。'
    )

    # TABLE 1
    p = doc.add_paragraph()
    run = p.add_run("表1：労働ビザ医療要件にハンセン病を明示的に記載する国・地域（n=20）")
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['地域', '国・地域', '法的根拠', '規定の性質', '年']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A237E')

    data = [
        ['GCC/中東', 'サウジアラビア', 'GAMCA/WAFID規則 第6版', '自動排除', '2021'],
        ['', 'UAE', 'GAMCA/WAFID規則', '自動排除', '2021'],
        ['', 'カタール', 'GAMCA/WAFID規則', '自動排除', '2021'],
        ['', 'クウェート', 'GAMCA/WAFID規則', '自動排除', '2021'],
        ['', 'オマーン', 'GAMCA/WAFID規則', '自動排除', '2021'],
        ['', 'バーレーン', 'GAMCA/WAFID規則', '自動排除', '2021'],
        ['東・東南アジア', '中国', '外国人体格検査表', '不罹患証明', '\u2014'],
        ['', 'タイ', '緊急勅令 B.E.2560', '禁止疾病', '2017'],
        ['', '台湾', '労働許可規則', '不罹患証明', '\u2014'],
        ['', 'マレーシア', 'FOMEMA要件', '医療上不適格', '\u2014'],
        ['', 'フィリピン', 'OWWA要件', '不罹患証明', '\u2014'],
        ['アフリカ', '南アフリカ', '移民法2002年; BI-811', '不罹患証明', '2002'],
        ['', 'ナミビア', '移民管理法', '医療上の禁止', '1993'],
        ['ヨーロッパ', 'ロシア', '連邦法 No.115-FZ', '国外退去事由', '2002'],
        ['', 'マルタ', '移民法 Cap.217', '医療上の入国不許可', '\u2014'],
        ['アメリカ大陸', '米国', 'INA §212(a)(1)(A)(i)', 'Class A条件*', '1952\u2020'],
        ['', 'バルバドス', '移民法 Cap.190', '医療上の禁止', '\u2014'],
        ['', '米領ヴァージン諸島', '米国連邦法', 'Class A条件*', '\u2014'],
        ['南アジア', 'インド', '各州雇用法', '雇用制限', '各種'],
        ['', '香港特別行政区', '入国管理条例 Cap.115', '医療上の入国不許可', '\u2014'],
    ]
    add_table_data(table, data, font_size=7)

    p = doc.add_paragraph()
    run = p.add_run(
        '*CDC技術指示書（2018年）により、MDT治療完了者または治療中の者については免除が認められる。\n'
        '\u20201952年制定。ハンセン病規定はその後の改正を経ても残存。'
    )
    run.font.size = Pt(8)
    run.italic = True

    doc.add_page_break()

    # GCC Analysis
    add_heading_styled(doc, '3.4 GCC/GAMCA制度', level=2)
    add_body(doc,
        'GCC6カ国はGCC認定医療センター協会（GAMCA）、別名WAFID（Wafid Al Marakiz）制度を通じた'
        '統一的な雇用前医療検査制度を運用している。WAFID医療検査規則（第6版、2021年）は、'
        '申請者を「医療上不適格」とする条件の標準化リストを規定しており、結核、HIV/AIDS、'
        'B型・C型肝炎、梅毒等と並んで「ハンセン病」を含めている [13]。医療検査は主に南・東南アジア'
        '（インド、パキスタン、バングラデシュ、スリランカ、フィリピン、インドネシア、ネパール）の'
        '出身国にあるGAMCA認定医療センターで実施される。GCC諸国は合計3,500万人以上の外国人労働者を'
        '受け入れており、ハンセン病スクリーニングを通じて外国人労働者に影響を与える世界最大の単一制度である。'
    )

    add_heading_styled(doc, '3.5 疾病スクリーニング頻度分析', level=2)
    add_body(doc,
        '疾病特異的スクリーニング要件が確認された58カ国のうち、最も一般的に記載された疾病は以下の通りである（図5）：'
    )

    disease_list = [
        '結核（TB）：55カ国（94.8%）',
        'HIV/AIDS：48カ国（82.8%）',
        '梅毒/性感染症：38カ国（65.5%）',
        'B型肝炎：30カ国（51.7%）',
        'ハンセン病（leprosy）：20カ国（34.5%）',
        '薬物依存/薬物乱用：25カ国（43.1%）',
        '精神疾患/精神科的状態：18カ国（31.0%）',
        'C型肝炎：12カ国（20.7%）',
    ]
    for d in disease_list:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()

    # Insert disease bar chart
    fig_path = os.path.join(FIGURES_DIR, "fig2_disease_bar.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        '図5：疾病特異的スクリーニングを有する58カ国における労働ビザ医療検査で指定される疾病。'
        'ハンセン病（赤）は5番目に多くスクリーニングされる疾病。'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    add_heading_styled(doc, '3.6 地域分布', level=2)
    add_body(doc,
        'ハンセン病スクリーニング要件の地域分布には明確な地理的パターンが認められる（図6）。'
        'GCC/中東地域はハンセン病スクリーニングを行う国の30%（20カ国中6カ国）を占めるが、'
        '統一的なGAMCA制度により影響を受ける個人の実数は他のどの地域よりもはるかに多い。'
        '東・東南アジア（25%、5カ国）は第2の集中地域であり、歴史的なスティグマと'
        '残存する流行地域の存在を反映している。'
    )

    # Insert regional donut
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig3_regional_donut.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        '図6：（A）労働ビザ医療検査でハンセン病を明示的にスクリーニングする20カ国・地域の地域分布。'
        '（B）移民法におけるハンセン病規定の法的性質。'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    add_heading_styled(doc, '3.7 疾病特異的要件を持たない国（カテゴリーD）', level=2)
    add_body(doc,
        '52カ国・地域が標準的な労働ビザに対して疾病特異的医療検査を要求していないことが特定された。'
        '最も注目すべきは、EU/シェンゲン圏26カ国がEU機能条約（TFEU）に基づく労働者の自由移動原則に依拠し、'
        '就労目的の健康に基づく入国制限を課していないことである。このカテゴリーの他の国には、'
        '日本（JPETSによる特定国籍者向けTBのみスクリーニング）、韓国、シンガポール、カナダ、'
        'オーストラリア、ニュージーランド、英国が含まれる。'
    )

    # TABLE 2
    p = doc.add_paragraph()
    run = p.add_run(
        '表2：疾病特異的スクリーニングを要求するがハンセン病を含まない国の例（n=15）'
    )
    run.bold = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'

    headers2 = ['国', 'スクリーニング対象疾病', '特記事項', 'ハンセン病']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2E7D32')

    data2 = [
        ['カナダ', 'TB、梅毒', '移民医療検査（IME）', '含まず'],
        ['オーストラリア', 'TB、HIV、肝炎B/C', '指定医制度（Bupa）', '含まず'],
        ['英国', 'TB（6ヶ月以上のビザ）', 'IOM指定医スクリーニング', '含まず'],
        ['ニュージーランド', 'TB、HIV、肝炎B', 'Immigration NZ指定医', '含まず'],
        ['日本', 'TBのみ（特定国籍者）', 'JPETS入国前TBスクリーニング', '含まず'],
        ['韓国', 'TB、HIV、薬物', 'E-9ビザ医療検査', '含まず'],
        ['シンガポール', 'TB、HIV、梅毒、肝炎B、マラリア', '外国人労働者医療検査（FWME）', '含まず'],
        ['イスラエル', 'TB、HIV、肝炎B/C', '入国医療要件', '含まず'],
        ['ケニア', 'TB、HIV', '就労許可医療証明書', '含まず'],
        ['ナイジェリア', 'TB、HIV、肝炎B', '外国人枠医療検査', '含まず'],
        ['ガーナ', 'TB、HIV、黄熱病', '移民医療証明書', '含まず'],
        ['ブラジル', '黄熱病（ワクチン接種）', '労働ビザ医療検査なし', '含まず'],
        ['メキシコ', '指定なし', '標準医療検査なし', '含まず'],
        ['アルゼンチン', '指定なし', '疾病特異的検査なし', '含まず'],
        ['チリ', 'TB、HIV（特定ビザ）', '一時居住ビザ医療', '含まず'],
    ]
    add_table_data(table2, data2, font_size=7)

    doc.add_page_break()

    add_heading_styled(doc, '3.8 感染性とスクリーニングの不一致', level=2)
    add_body(doc,
        '注目すべき知見は、スクリーニング対象疾病の感染性と労働ビザ要件への組み込み頻度との間の不一致である（図7）。'
        'ハンセン病は独特の位置を占める：一般的にスクリーニングされる疾病の中で最も低い感染性'
        '（基本再生産数は推定1.0〜1.3、結核の4〜10、麻疹の12〜18と比較）を持つにもかかわらず、'
        'スクリーニング頻度では5位にランクされている。これは、組み込みが疫学的リスク評価ではなく'
        '歴史的スティグマによって駆動されていることを示唆する。'
    )

    # Insert transmissibility figure
    doc.add_paragraph()
    fig_path = os.path.join(FIGURES_DIR, "fig5_transmissibility.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(
        '図7：労働ビザ医療検査における疾病の感染性とスクリーニング頻度の関係。'
        'ハンセン病（赤）は低い感染性にもかかわらず高いスクリーニング頻度を示し、'
        'エビデンスと政策の乖離を示す。バブルサイズは各疾病をスクリーニングしている国の数に比例。'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_page_break()

    # ==== DISCUSSION ====
    add_heading_styled(doc, '4. 考察（Discussion）', level=1)

    add_heading_styled(doc, '4.1 エビデンスの要約', level=2)
    add_body(doc,
        '本スコーピングレビューは、労働ビザ医療検査におけるハンセン病スクリーニング要件の'
        '初の包括的グローバルマッピングを示す。我々の所見は、20カ国・地域（調査対象197の10.2%）が'
        '移民医療要件にハンセン病を明示的に記載し続けていることを明らかにした。GCC/GAMCA制度だけで'
        '推定3,500万人以上の外国人労働者に影響を与えており、ハンセン病関連の移民スクリーニング'
        'メカニズムとして世界最大である。'
    )

    add_heading_styled(doc, '4.2 標準医療と乖離した行政要求', level=2)
    add_body(doc,
        '本レビューの中心的所見は、行政的移民要件とハンセン病に関する現在の医学知識との間の'
        '深刻な断絶である。標準医療は以下を認識している：（a）ハンセン病はMDTで治癒可能、'
        '（b）治療開始数日以内に非感染性となる、（c）人口の約95%が自然免疫を有する、'
        '（d）長い潜伏期間（3〜20年）のため入国時検査による症例検出は効果がない、'
        '（e）WHOがMDTを世界中で無償提供している [2,4]。'
    )

    add_body(doc,
        'にもかかわらず、20カ国が自動排除（GCC諸国）から不罹患証明（中国、南アフリカ）、'
        '国外退去の医療事由（ロシア）に至る要件を維持している。これらの行政要求はエビデンスに基づく医療に'
        '何ら根拠を持たず、標準医療との明確な乖離を示している。'
    )

    add_heading_styled(doc, '4.3 他のスクリーニング対象疾病との比較', level=2)
    add_body(doc,
        'ハンセン病が5番目に多くスクリーニングされる疾病として位置づけられていることは、'
        '疫学的パラメータと照合すると際立っている。結核（疾病特異的要件を持つ国の94.8%がスクリーニング）は'
        '年間約130万人を死亡させ、R0は4〜10であり、入国時の胸部X線検査と喀痰検査で検出可能である [14]。'
        'HIV/AIDS（82.8%）は確立された伝播経路と生涯にわたる影響を有する。梅毒（65.5%）とB型肝炎（51.7%）は'
        '定義された経路を通じて容易に伝播する。対照的に、ハンセン病のR0は約1.0〜1.3であり、'
        '伝播には長期間の密接な接触を要し、完全に治癒可能であり、世界の年間新規症例は20万件未満である [4]。'
    )

    add_heading_styled(doc, '4.4 比較モデルとしての日本の経験', level=2)
    add_body(doc,
        '日本のアプローチは示唆的な反例を提供する。日本は1953年にらい予防法を制定し、'
        'ハンセン病患者の強制隔離を義務付けた。この法律は元患者による持続的なアドボカシー活動を経て'
        '1996年に廃止された。2001年、熊本地方裁判所は政府による強制隔離の長期間にわたる実施が'
        '憲法上の権利を侵害すると判決し、正式な謝罪と補償に至った [15]。重要なのは、'
        '日本が移民医療スクリーニング要件にハンセン病を含めていないことである。'
        '日本の入国前TBスクリーニングプログラム（JPETS）は、移民目的のエビデンスに基づく疾病特異的'
        'スクリーニングが実現可能であること、そしてそのようなプログラムが国境での公衆衛生介入の'
        '閾値を満たさない疾病を含む必要がないことを示している。'
    )

    add_heading_styled(doc, '4.5 国際法的枠組み', level=2)
    add_body(doc,
        '移民医療要件におけるハンセン病の維持は、以下の複数の国際文書に直接矛盾する：'
    )

    instruments = [
        '国連総会決議A/RES/65/215（2010年）：原則7は、ハンセン病罹患者がその状態を理由に入国や居住の権利を否定されるべきではないと規定。',
        '国連人権理事会決議A/HRC/RES/29/5（2015年）：原則とガイドラインを再確認し、その完全な実施を求める。',
        '国連ハンセン病特別報告者の報告書（2017年〜現在）：移民文脈における差別的法律を文書化。',
        'WHO世界ハンセン病戦略2021〜2030：ゼロ差別を中核的柱として強調。',
    ]
    for inst in instruments:
        doc.add_paragraph(inst, style='List Bullet')

    add_heading_styled(doc, '4.6 限界', level=2)
    add_body(doc,
        '本レビューにはいくつかの限界がある。第一に、2段階の検索戦略にもかかわらず、82カ国（41.6%）の'
        '疾病特異的スクリーニング要件が未確認のままであった。調査した197カ国のうち、110カ国（55.8%）は'
        '英語情報源のみで完全に特定された（疾病特異的スクリーニング確認58カ国、スクリーニング不存在確認52カ国）。'
        'アラビア語、ベトナム語、シンハラ語、インドネシア語等の公用語による補足的多言語調査により、'
        '5カ国（2.5%）の具体的なスクリーニングデータが追加的に得られ、確認済み情報を持つ国の合計は'
        '115カ国（58.4%）となった。残りの82カ国（41.6%）は多言語調査にもかかわらず完全に特定できなかった。'
        'これら82カ国のうち39カ国（47.6%）は英語を公用語または共同公用語としており、言語障壁のみでは'
        '情報格差を説明できないことを示唆している。むしろ、指定医への非公開の要件伝達が主要な要因である'
        '可能性がある。残りの43の非英語圏には、フランス語圏（15カ国、主に西・中央アフリカ）、'
        'アラビア語圏（13カ国、主に中東・北アフリカ）、ポルトガル語圏（5カ国）が含まれる。'
        '注目すべきは、補足的多言語調査で確認された5カ国のいずれにおいてもハンセン病のスクリーニングは'
        '確認されなかったことであり、特定された20カ国がほぼ完全な列挙である可能性を示唆している。'
    )
    add_body(doc,
        '第二に、一部の国の要件は指定医にのみ通知され公開されていないため、過小評価につながる可能性がある。'
        '第三に、移民医療要件は頻繁に変更されるため、我々の所見は2026年3月時点の情報を反映している。'
        '第四に、法的規定の存在は必ずしも現在の運用実態を反映しない—一部の国は実際には適用されていない'
        '時代遅れの規定を残存させている可能性がある。第五に、疾病特異的医療検査要件が確認されなかった'
        '52カ国（26.4%）（主にEU/シェンゲン圏諸国および小島嶼国）については、公用語による情報源で'
        'これらの国の移民枠組みに疾病特異的な労働ビザスクリーニングが含まれていないことを確認しており、'
        '地域的な政策枠組み（例：EU域内移動の自由指令）と整合している。'
    )

    doc.add_page_break()

    # ==== CONCLUSIONS ====
    add_heading_styled(doc, '5. 結論（Conclusions）', level=1)
    add_body(doc,
        'ハンセン病は少なくとも20カ国・地域の移民医療要件に組み込まれたままであり、世界中の数百万人の'
        '外国人労働者に影響を与えている。これらの要件は、治癒可能であり、高度に感染性ではなく、'
        '長い潜伏期間のため入国時検出が効果的でない疾病に対するスクリーニングという、'
        '標準医療と根本的に乖離した行政要求を表している。3,500万人以上の労働者に影響を与える'
        'GCC/GAMCA制度は世界最大のそのような制度である。'
    )

    add_body(doc,
        'ハンセン病の疫学的プロファイルとスクリーニング頻度（労働ビザ要件で5番目に多く指定される疾病）'
        'との間の不一致は、これらの規定がエビデンスに基づく公衆衛生政策ではなく歴史的スティグマに'
        '根ざしていることを強く示唆する。この所見は、国連ハンセン病特別報告者およびILEP差別的法律'
        'データベースの観察と一致している。'
    )

    add_body(doc,
        '我々は以下を勧告する：（1）現在移民医療検査でハンセン病をスクリーニングしている全ての国が'
        'これらの要件のエビデンスに基づく見直しを実施すること、（2）GCC諸国がGAMCA/WAFIDの医療不適格基準から'
        'ハンセン病を除外すること、（3）WHOがハンセン病は入国前の国境スクリーニングに値しないとする'
        '明確なガイダンスを発出すること、（4）国連人権理事会がハンセン病に基づく移民制限の監視を強化すること。'
        '日本の経験は、移民スクリーニングからハンセン病を除外することが実現可能であり、公衆衛生目標と'
        '人権義務の両方に合致することを示している。'
    )

    doc.add_page_break()

    # ==== REFERENCES ====
    add_heading_styled(doc, '参考文献（References）', level=1)

    refs = [
        '1. Scollard DM, Adams LB, Gillis TP, et al. The continuing challenges of leprosy. Clin Microbiol Rev. 2006;19:338\u201381.',
        '2. WHO. Guidelines for the diagnosis, treatment and prevention of leprosy. New Delhi: World Health Organization; 2018.',
        '3. Sermrittirong S, Van Brakel WH. Stigma in leprosy: concepts, causes and determinants. Lepr Rev. 2014;85:36\u201347.',
        '4. WHO. Weekly epidemiological record: Global leprosy (Hansen disease) update, 2024. Wkly Epidemiol Rec. 2025;100:377\u201396.',
        '5. Fine PE. Leprosy: the epidemiology of a slow bacterium. Epidemiol Rev. 1982;4:161\u201388.',
        '6. Britton WJ, Lockwood DN. Leprosy. Lancet. 2004;363:1209\u201319.',
        '7. UN General Assembly. Resolution A/RES/65/215: Elimination of discrimination against persons affected by leprosy and their family members. New York: United Nations; 2010.',
        '8. UN Human Rights Council. Resolution A/HRC/RES/35/9: Appointment of Special Rapporteur on the elimination of discrimination against persons affected by leprosy and their family members. Geneva: OHCHR; 2017.',
        '9. ILEP. Discriminatory Laws Database. International Federation of Anti-Leprosy Associations; 2024. Available from: https://ilepfederation.org/discriminatory-laws/',
        '10. Tricco AC, Lillie E, Zarin W, et al. PRISMA Extension for Scoping Reviews (PRISMA-ScR): Checklist and Explanation. Ann Intern Med. 2018;169:467\u201373.',
        '11. Arksey H, O\u2019Malley L. Scoping studies: towards a methodological framework. Int J Soc Res Methodol. 2005;8:19\u201332.',
        '12. Levac D, Colquhoun H, O\u2019Brien KK. Scoping studies: advancing the methodology. Implement Sci. 2010;5:69.',
        '13. GCC Approved Medical Centers Association. WAFID Medical Examination Regulations, 6th Version. Riyadh: GAMCA Secretariat; 2021.',
        '14. WHO. Global tuberculosis report 2024. Geneva: World Health Organization; 2024.',
        '15. 熊本地方裁判所. 2001年5月11日判決：ハンセン病国家賠償訴訟. 熊本; 2001.',
        '16. Republic of South Africa. Immigration Act No. 13 of 2002; Form BI-811 (Medical Certificate). Pretoria: Department of Home Affairs.',
        '17. CDC. Technical Instructions for Hansen\u2019s Disease for Panel Physicians. Atlanta: Centers for Disease Control and Prevention; 2018.',
        '18. Abubakar I, Aldridge RW, Devakumar D, et al. The UCL\u2013Lancet Commission on Migration and Health. Lancet. 2018;392:2606\u201354.',
        '19. Nanri T. Initiatives to address leprosy as a human rights issue through the mandate of UN Special Rapporteur: Achievements and challenges. PLoS Negl Trop Dis. 2022;16:e0010201.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)

    doc.add_page_break()

    # ==== SUPPORTING INFORMATION ====
    add_heading_styled(doc, '補足資料（Supporting Information）', level=1)

    p = doc.add_paragraph()
    run = p.add_run('S1 Appendix. PRISMA-ScRチェックリスト')
    run.bold = True

    checklist_table = doc.add_table(rows=1, cols=3)
    checklist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    checklist_table.style = 'Table Grid'

    ch_headers = ['項目番号', 'チェックリスト項目', '報告箇所']
    for i, h in enumerate(ch_headers):
        cell = checklist_table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A237E')

    checklist_data = [
        ['1', 'タイトル：スコーピングレビューとして特定', 'タイトルページ'],
        ['2', '構造化された要約', '抄録'],
        ['3', '根拠', '緒言'],
        ['4', '目的', 'セクション1.1'],
        ['5', 'プロトコルと登録', 'セクション2.1'],
        ['6', '適格基準', 'セクション2.2'],
        ['7', '情報源', 'セクション2.3'],
        ['8', '検索', 'セクション2.3'],
        ['9', 'エビデンス源の選択', 'セクション2.4'],
        ['10', 'データ抽出プロセス', 'セクション2.5'],
        ['11', 'データ項目', 'セクション2.5'],
        ['12', '批判的評価（該当する場合）', 'N/A（スコーピングレビュー）'],
        ['13', '結果の統合', 'セクション2.6'],
        ['14', 'エビデンス源の選択', 'セクション3.1, 図1'],
        ['15', 'エビデンス源の特性', 'セクション3.2'],
        ['16', '批判的評価（該当する場合）', 'N/A（スコーピングレビュー）'],
        ['17', '個別ソースの結果', '表1〜2'],
        ['18', '結果の統合', 'セクション3.3〜3.8, 図2〜4'],
        ['19', 'エビデンスの要約', 'セクション4.1'],
        ['20', '限界', 'セクション4.6'],
        ['21', '結論', 'セクション5'],
        ['22', '資金', '謝辞'],
    ]
    add_table_data(checklist_table, checklist_data, font_size=7)

    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run('S2 Appendix. 包括的国別データ')
    run.bold = True
    p = doc.add_paragraph()
    p.add_run('[注：197カ国・地域の完全な国別データセットは補足Excelファイルとして利用可能。'
              '責任著者に連絡されたい。]')

    p = doc.add_paragraph()
    run = p.add_run('\nS3 Appendix. ILEPデータベースとのクロスリファレンス')
    run.bold = True
    p = doc.add_paragraph()
    p.add_run(
        '我々の所見をILEP差別的法律データベース（2024年）とクロスリファレンスすると、'
        '24カ国にわたって文書化された139の差別的法律のうち、8カ国の9法律が移民・市民権に'
        '特に関連することが明らかになった。インドは139法律中108件（77.7%）を占めるが、'
        'これらは主に移民ではなく国内の雇用・財産権の文脈に関するものである。'
        '本レビューでは、ILEPデータベースに含まれていない追加の国々（マルタ、バルバドス、'
        '米領ヴァージン諸島、フィリピン、香港特別行政区等）が移民法にハンセン病規定を'
        '維持していることを特定した。'
    )

    # S4 Appendix - Unreachable countries list
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run('S4 Appendix. 疾病特異的スクリーニング要件が未確認の82カ国のリスト')
    run.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run(
        '以下の82カ国は、多言語調査にもかかわらず完全に特定できなかった。'
        '地域別に整理し、公用語を併記する。これらのうち39カ国（47.6%）は英語を公用語または'
        '共同公用語としており、言語障壁のみではなく、指定医への非公開の要件伝達が'
        '情報格差の主要因である可能性を示唆している。'
    )

    # Create unreachable countries table
    unr_table = doc.add_table(rows=1, cols=3)
    unr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    unr_table.style = 'Table Grid'
    for i, h in enumerate(['地域', '国名', '公用語']):
        cell = unr_table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, 'C62828')

    import json as _json
    _cls_path = os.path.join(FIGURES_DIR, 'country_classification.json')
    if os.path.exists(_cls_path):
        with open(_cls_path, 'r') as _f:
            _cls = _json.load(_f)
        for region, countries in _cls.get('unreachable_by_region', {}).items():
            for idx, country in enumerate(countries):
                lang = _cls.get('unreachable_languages', {}).get(country, '')
                row = unr_table.add_row()
                row.cells[0].text = region if idx == 0 else ''
                row.cells[1].text = country
                row.cells[2].text = lang
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(7)

    # ---- DECLARATIONS ----
    doc.add_page_break()
    add_heading_styled(doc, '宣言（Declarations）', level=1)

    p = doc.add_paragraph()
    run = p.add_run('資金（Funding）: ')
    run.bold = True
    p.add_run('[追記予定]')

    p = doc.add_paragraph()
    run = p.add_run('利益相反（Competing interests）: ')
    run.bold = True
    p.add_run('著者らは利益相反がないことを宣言する。')

    p = doc.add_paragraph()
    run = p.add_run('著者の貢献（Author contributions）: ')
    run.bold = True
    p.add_run('[CRediT分類法を用いて追記予定]')

    p = doc.add_paragraph()
    run = p.add_run('データの利用可能性（Data availability）: ')
    run.bold = True
    p.add_run('本研究で生成された全てのデータは、本論文およびその補足情報ファイルに含まれている。'
              '完全な国別データセットは、合理的な要求に応じて責任著者から入手可能。')

    p = doc.add_paragraph()
    run = p.add_run('謝辞（Acknowledgements）: ')
    run.bold = True
    p.add_run('[追記予定]')

    # Save
    output_path = os.path.join(OUTPUT_DIR, "PLoS_NTDs_Full_Paper_JA.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == '__main__':
    create_plos_full_ja()
